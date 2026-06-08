from pathlib import Path
from zipfile import ZipFile

from docx import Document

from reporting.stage9_bmc import (
    TITLE,
    _abstract_text,
    _document_text,
    _load_scrna_support,
    bmc_paths,
    language_risk_report,
    load_stage8_evidence,
)


ROOT = Path(__file__).resolve().parents[1]


def test_bmc_abstract_is_structured_and_bounded():
    evidence = load_stage8_evidence(ROOT)
    scrna = _load_scrna_support(ROOT)
    abstract = _abstract_text(evidence, scrna)
    assert list(abstract) == ["Background", "Methods", "Results", "Conclusions"]
    assert sum(len(text.split()) for text in abstract.values()) <= 350
    assert "208,506" in abstract["Methods"]
    assert "26" in abstract["Results"]


def test_bmc_documents_exist_and_keep_claim_boundaries():
    paths = bmc_paths(ROOT)
    for key in ("main", "title_page", "cover_letter", "declarations"):
        assert paths[key].exists()
        assert paths[key].stat().st_size > 10_000
    text = _document_text(paths["main"])
    assert TITLE in text
    assert "qualitative/IHC-link evidence" in text
    assert "exploratory quantitative proteomic support" in text
    assert "208,506 cells" in text
    assert "broader cycling-cell signal" in text
    assert "did not show stable improvement" in text


def test_bmc_main_docx_has_line_and_page_numbering():
    path = bmc_paths(ROOT)["main"]
    with ZipFile(path) as archive:
        xml = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
    assert 'w:restart="continuous"' in xml
    assert " PAGE " in xml
    assert 'w:line="480"' in xml or 'w:lineRule="auto"' in xml


def test_bmc_main_docx_has_ordered_assets():
    document = Document(bmc_paths(ROOT)["main"])
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Fig. 1" in text
    assert "Fig. 5" in text
    assert "Supplementary Fig. S1" in text
    assert "Supplementary Fig. S7" in text
    assert "Supplementary Table S1" in text
    assert "Supplementary Table S8" in text
    assert len(document.tables) == 4


def test_bmc_language_audit_has_no_prohibited_positive_claims():
    report = language_risk_report(ROOT)
    assert "High-risk phrase hits: 0" in report
