from pathlib import Path

from docx import Document

from reporting.stage10a_bmc_cleanup import (
    stage10a_paths,
    validate_cleaned_manuscript,
)


ROOT = Path(__file__).resolve().parents[1]


def test_stage10a_outputs_exist():
    paths = stage10a_paths(ROOT)
    for key in (
        "cleaned",
        "reference_report",
        "table_report",
        "figure_report",
        "author_queries",
    ):
        assert paths[key].exists()
        assert paths[key].stat().st_size > 500


def test_stage10a_cleaned_doc_has_no_placeholders_or_local_paths():
    checks = validate_cleaned_manuscript(ROOT)
    assert checks["author_placeholder"] == 0
    assert checks["confirmation_placeholder"] == 0
    assert checks["reference_placeholder"] == 0
    assert checks["local_windows_path"] == 0
    assert checks["author_present"]
    assert checks["funding_present"]
    document = Document(stage10a_paths(ROOT)["cleaned"])
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "国家优秀青年医师" in text
    assert "https://github.com/HAOYANGLI888/SC-PROST-LUAD" in text
    assert "additional institutional ethics approval was not sought" in text.lower()


def test_stage10a_reference_completion():
    paths = stage10a_paths(ROOT)
    document = Document(paths["cleaned"])
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for doi in (
        "10.1038/s41467-020-16164-1",
        "10.1158/0008-5472.CAN-11-1403",
        "10.1097/JTO.0000000000000042",
        "10.1038/onc.2015.375",
        "10.1038/nm.1790",
        "10.1186/s12916-014-0241-z",
        "10.7326/M18-1376",
        "10.1056/NEJMp1607591",
        "10.1002/sim.4154",
        "10.1186/s12874-018-0482-1",
        "10.1038/s41591-018-0096-5",
        "10.1158/2159-8290.CD-20-1285",
        "10.1016/j.cell.2020.07.017",
        "10.1038/s41586-024-07113-9",
        "10.1186/s12916-019-1466-7",
    ):
        assert doi.lower() in text.lower()
    references = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.startswith(tuple(f"{index}. " for index in range(1, 25)))
    ]
    assert len(references) == 24


def test_stage10a_tables_are_simplified_and_landscape():
    document = Document(stage10a_paths(ROOT)["cleaned"])
    assert len(document.tables) == 4
    assert [cell.text for cell in document.tables[2].rows[0].cells] == [
        "Program",
        "Bulk risk association",
        "GEO consistency",
        "Raw scRNA localization",
        "Final interpretation",
    ]
    assert [cell.text for cell in document.tables[3].rows[0].cells] == [
        "Evidence layer",
        "Coverage",
        "Main finding",
        "Boundary",
    ]
    assert any(section.page_width > section.page_height for section in document.sections)


def test_stage10a_citation_order_and_numbering_pass():
    checks = validate_cleaned_manuscript(ROOT)
    assert checks["figure_order"]
    assert checks["table_order"]
    assert checks["supp_figure_order"]
    assert checks["supp_table_order"]
    assert checks["line_numbering"]
    assert checks["page_numbering"]
