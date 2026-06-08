from pathlib import Path
from zipfile import ZipFile

from docx import Document

from reporting.stage9_jtm import (
    TITLE,
    _abstract_text,
    _document_text,
    language_risk_report,
    load_stage8_evidence,
    stage9_paths,
)


ROOT = Path(__file__).resolve().parents[1]


def test_stage9_abstract_is_structured_and_bounded():
    evidence = load_stage8_evidence(ROOT)
    abstract = _abstract_text(evidence)
    assert list(abstract) == ["Background", "Methods", "Results", "Conclusions"]
    word_count = sum(len(text.split()) for text in abstract.values())
    assert word_count <= 350
    assert "26" in abstract["Results"]
    assert "qualitative" in abstract["Results"]


def test_stage9_documents_exist_and_keep_required_boundaries():
    paths = stage9_paths(ROOT)
    for key in ("main", "title_page", "cover_letter", "declarations"):
        assert paths[key].exists()
        assert paths[key].stat().st_size > 10_000
    text = _document_text(paths["main"])
    assert TITLE in text
    assert "qualitative/IHC-link evidence" in text
    assert "exploratory quantitative proteomic support" in text
    assert "curated signatures" in text
    assert "did not show stable improvement" in text
    assert "modest" in text


def test_stage9_main_docx_has_line_and_page_numbering():
    path = stage9_paths(ROOT)["main"]
    with ZipFile(path) as archive:
        xml = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
    assert 'w:restart="continuous"' in xml
    assert " PAGE " in xml
    assert 'w:line="480"' in xml or 'w:lineRule="auto"' in xml


def test_stage9_main_docx_has_ordered_tables_and_figure_legends():
    document = Document(stage9_paths(ROOT)["main"])
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Fig. 1" in text
    assert "Fig. 5" in text
    assert "Additional file 1" in text
    assert "Additional file 7" in text
    assert len(document.tables) == 4


def test_stage9_language_audit_has_no_positive_overclaim_hits():
    report = language_risk_report(ROOT)
    assert "High-risk phrase hits: 0" in report
