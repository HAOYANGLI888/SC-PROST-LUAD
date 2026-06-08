"""Journal of Translational Medicine manuscript generation and audit utilities."""

from __future__ import annotations

from datetime import date, datetime
from pathlib import Path
import re
from typing import Iterable
from zipfile import ZipFile

import pandas as pd

from reporting.stage8 import Stage8Evidence, load_stage8_evidence


TITLE = (
    "An externally validated transcriptomic-clinical survival model with "
    "cell-state and proteomic interpretation in lung adenocarcinoma"
)
RUNNING_TITLE = "Externally validated LUAD survival model"
JTM_GUIDELINES_URL = "https://link.springer.com/journal/12967/submission-guidelines"
JTM_CONDITIONS_URL = (
    "https://translational-medicine.biomedcentral.com/submission-guidelines/"
    "conditions-of-publication"
)


class Stage9Error(RuntimeError):
    """Raised when Stage 9 cannot generate an evidence-grounded artifact."""


def stage9_paths(root: str | Path = ".") -> dict[str, Path]:
    project_root = Path(root).resolve()
    return {
        "root": project_root,
        "manuscript": project_root / "outputs" / "manuscript",
        "tables": project_root / "outputs" / "tables",
        "reports": project_root / "outputs" / "reports",
        "main": project_root
        / "outputs"
        / "manuscript"
        / "JTM_main_manuscript_draft.docx",
        "title_page": project_root
        / "outputs"
        / "manuscript"
        / "JTM_title_page.docx",
        "cover_letter": project_root
        / "outputs"
        / "manuscript"
        / "JTM_cover_letter.docx",
        "declarations": project_root
        / "outputs"
        / "manuscript"
        / "JTM_declarations.docx",
        "checklist": project_root
        / "outputs"
        / "tables"
        / "JTM_submission_file_checklist.csv",
        "compliance": project_root
        / "outputs"
        / "reports"
        / "JTM_compliance_report.md",
        "language": project_root
        / "outputs"
        / "reports"
        / "JTM_language_risk_audit.md",
        "claims": project_root
        / "outputs"
        / "reports"
        / "JTM_evidence_claim_audit.md",
        "audit": project_root / "audit_report.md",
    }


def _require_docx():
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION_START
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise Stage9Error(
            "python-docx is required. Install it with "
            "`python -m pip install python-docx>=1.1`."
        ) from exc
    return {
        "Document": Document,
        "WD_SECTION_START": WD_SECTION_START,
        "WD_CELL_VERTICAL_ALIGNMENT": WD_CELL_VERTICAL_ALIGNMENT,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "OxmlElement": OxmlElement,
        "qn": qn,
        "Inches": Inches,
        "Pt": Pt,
    }


def _set_cell_shading(cell, fill: str) -> None:
    docx = _require_docx()
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = docx["OxmlElement"]("w:shd")
    shading.set(docx["qn"]("w:fill"), fill)
    tc_pr.append(shading)


def _add_page_number(paragraph) -> None:
    docx = _require_docx()
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = docx["OxmlElement"]("w:fldChar")
    begin.set(docx["qn"]("w:fldCharType"), "begin")
    instruction = docx["OxmlElement"]("w:instrText")
    instruction.set(docx["qn"]("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = docx["OxmlElement"]("w:fldChar")
    separate.set(docx["qn"]("w:fldCharType"), "separate")
    text = docx["OxmlElement"]("w:t")
    text.text = "1"
    end = docx["OxmlElement"]("w:fldChar")
    end.set(docx["qn"]("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def _enable_continuous_line_numbers(section) -> None:
    docx = _require_docx()
    sect_pr = section._sectPr
    for existing in sect_pr.findall(docx["qn"]("w:lnNumType")):
        sect_pr.remove(existing)
    line_numbers = docx["OxmlElement"]("w:lnNumType")
    line_numbers.set(docx["qn"]("w:countBy"), "1")
    line_numbers.set(docx["qn"]("w:start"), "1")
    line_numbers.set(docx["qn"]("w:restart"), "continuous")
    line_numbers.set(docx["qn"]("w:distance"), "360")
    sect_pr.append(line_numbers)


def _new_document(*, letter: bool = False):
    docx = _require_docx()
    document = docx["Document"]()
    section = document.sections[0]
    section.top_margin = docx["Inches"](1.0)
    section.bottom_margin = docx["Inches"](1.0)
    section.left_margin = docx["Inches"](1.0)
    section.right_margin = docx["Inches"](1.0)
    if letter:
        section.top_margin = docx["Inches"](0.8)
        section.bottom_margin = docx["Inches"](0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = docx["Pt"](12)
    normal._element.rPr.rFonts.set(docx["qn"]("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = docx["Pt"](0)

    for level in range(1, 4):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = docx["Pt"](14 if level == 1 else 12)
        style.font.bold = True
        style.font.color.rgb = None
        style._element.rPr.rFonts.set(docx["qn"]("w:eastAsia"), "Times New Roman")
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_before = docx["Pt"](6)
        style.paragraph_format.space_after = docx["Pt"](0)

    footer = section.footer.paragraphs[0]
    footer.alignment = docx["WD_ALIGN_PARAGRAPH"].CENTER
    _add_page_number(footer)
    _enable_continuous_line_numbers(section)
    document.core_properties.title = TITLE
    document.core_properties.subject = (
        "Journal of Translational Medicine submission draft"
    )
    document.core_properties.comments = (
        "Generated from locked project outputs; author metadata requires completion."
    )
    return document


def _add_title(document, text: str = TITLE) -> None:
    docx = _require_docx()
    paragraph = document.add_paragraph()
    paragraph.alignment = docx["WD_ALIGN_PARAGRAPH"].CENTER
    paragraph.paragraph_format.line_spacing = 2.0
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = docx["Pt"](16)


def _add_paragraph(
    document,
    text: str,
    *,
    bold_prefix: str | None = None,
    first_line: bool = True,
    style: str | None = None,
):
    docx = _require_docx()
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = docx["Pt"](0)
    if first_line and not style:
        paragraph.paragraph_format.first_line_indent = docx["Inches"](0.3)
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def _add_heading(document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def _add_table(document, title: str, columns: list[str], rows: Iterable[Iterable]):
    _add_paragraph(document, title, first_line=False)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        cell = table.rows[0].cells[index]
        cell.text = str(column)
        _set_cell_shading(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = _require_docx()["Pt"](9)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = "" if pd.isna(value) else str(value)
            cells[index].vertical_alignment = _require_docx()[
                "WD_CELL_VERTICAL_ALIGNMENT"
            ].CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = _require_docx()["Pt"](0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = _require_docx()["Pt"](9)
    return table


def _external_rows(evidence: Stage8Evidence) -> list[dict]:
    cox = evidence.external_cox.set_index("cohort")
    order = ["GSE31210", "GSE50081", "GSE72094", "GSE68465"]
    performance = evidence.external_performance.set_index("cohort")
    rows: list[dict] = []
    for cohort in order:
        row = performance.loc[cohort]
        cox_row = cox.loc[cohort]
        rows.append(
            {
                "cohort": cohort,
                "n": int(row["patient_count"]),
                "deaths": int(row["death_events"]),
                "c_index": float(row["c_index"]),
                "auc_1y": float(row["auc_1_year"]),
                "auc_3y": float(row["auc_3_year"]),
                "auc_5y": float(row["auc_5_year"]),
                "hr": float(cox_row["hazard_ratio_per_sd"]),
                "p": float(cox_row["p_value"]),
                "gene_missing": float(row["gene_missing_fraction"]),
            }
        )
    return rows


def _format_p(value: float) -> str:
    return f"{value:.3g}" if value >= 0.001 else f"{value:.2e}"


def _abstract_text(evidence: Stage8Evidence) -> dict[str, str]:
    external = ", ".join(
        f"{row['cohort']} {row['c_index']:.3f}" for row in _external_rows(evidence)
    )
    return {
        "Background": (
            "Transcriptomic survival models in lung adenocarcinoma often have "
            "limited external validation and uncertain biological interpretation "
            "across assay platforms."
        ),
        "Methods": (
            f"We developed an overall-survival model in {evidence.tcga_patients} "
            "TCGA-LUAD patients using clinical variables and a 25-component RNA "
            "representation. Preprocessing, feature construction and tuning were "
            "restricted to training folds. The frozen model was applied without "
            f"refitting to four GEO cohorts ({evidence.external_total} patients). "
            f"We then assessed {evidence.signature_count} curated cell-state "
            "signatures inferred from bulk RNA, qualitative Human Protein Atlas "
            "IHC links, exploratory CPTAC-LUAD proteomics and a separate WSI "
            "feasibility pilot."
        ),
        "Results": (
            "The RNA-plus-clinical Cox model achieved a nested-cross-validation "
            f"C-index of {evidence.nested_combined_cindex:.3f} "
            f"(SD {evidence.nested_combined_sd:.3f}), compared with "
            f"{evidence.nested_clinical_cindex:.3f} "
            f"(SD {evidence.nested_clinical_sd:.3f}) for clinical Cox. External "
            f"C-indices were {external}; the continuous risk score was associated "
            "with overall survival in all four cohorts, although the TCGA-derived "
            "cutoff was not portable. Higher risk was associated with hypoxia, "
            "proliferation, EMT-like tumor and CAF programs. HPA provided "
            f"qualitative or linkable evidence for {evidence.hpa_supported_count} "
            f"prespecified candidates. CPTAC quantified "
            f"{evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} "
            "candidates; LDHA, MKI67 and CDK1 showed direction-consistent nominal "
            "stage associations, and the hypoxia module was associated with "
            f"survival (HR {evidence.cptac_hypoxia_hr:.2f} per SD, 95% CI "
            f"{evidence.cptac_hypoxia_ci_low:.2f}-"
            f"{evidence.cptac_hypoxia_ci_high:.2f}, FDR "
            f"{evidence.cptac_hypoxia_fdr:.4f}). Only "
            f"{evidence.cptac_deaths} CPTAC deaths were available. The WSI pilot "
            "showed technical feasibility but no stable predictive gain."
        ),
        "Conclusions": (
            "A frozen transcriptomic-clinical score reproducibly ranked survival "
            "risk across four external cohorts and was associated with coherent "
            "tumor-state and microenvironment programs. Its modest incremental "
            "value, platform-dependent cutoff transfer and exploratory protein "
            "evidence warrant prospective, assay-harmonized evaluation before "
            "clinical use."
        ),
    }


def _build_main_tables(document, evidence: Stage8Evidence, root: Path) -> None:
    external = _external_rows(evidence)
    cohort_rows = [
        [
            "TCGA-LUAD",
            "RNA-seq",
            evidence.tcga_patients,
            evidence.tcga_deaths,
            evidence.tcga_censored,
            "Development/internal validation",
        ]
    ]
    cohort_rows.extend(
        [
            [
                row["cohort"],
                "Microarray",
                row["n"],
                row["deaths"],
                row["n"] - row["deaths"],
                "Locked external validation",
            ]
            for row in external
        ]
    )
    _add_table(
        document,
        "Table 1. Study cohorts used for model development and external validation.",
        ["Cohort", "Assay", "Patients", "Deaths", "Censored", "Role"],
        cohort_rows,
    )

    performance_rows = [
        [
            "TCGA nested CV",
            "Clinical Cox",
            f"{evidence.nested_clinical_cindex:.3f} ({evidence.nested_clinical_sd:.3f})",
            "-",
            "-",
            "-",
            "-",
        ],
        [
            "TCGA nested CV",
            "RNA PCA_25 + clinical Cox",
            f"{evidence.nested_combined_cindex:.3f} ({evidence.nested_combined_sd:.3f})",
            f"{evidence.nested_auc_1y:.3f}",
            f"{evidence.nested_auc_3y:.3f}",
            f"{evidence.nested_auc_5y:.3f}",
            "-",
        ],
    ]
    performance_rows.extend(
        [
            [
                row["cohort"],
                "Frozen full model",
                f"{row['c_index']:.3f}",
                f"{row['auc_1y']:.3f}",
                f"{row['auc_3y']:.3f}",
                f"{row['auc_5y']:.3f}",
                f"{row['hr']:.3f}; P={_format_p(row['p'])}",
            ]
            for row in external
        ]
    )
    _add_table(
        document,
        "Table 2. Discrimination and survival association of the frozen model.",
        ["Cohort", "Model", "C-index", "AUC 1 y", "AUC 3 y", "AUC 5 y", "HR per SD"],
        performance_rows,
    )

    correlations = pd.read_csv(
        root / "outputs" / "tables" / "stage4_tcga_cell_state_risk_correlation.csv"
    )
    consistency = pd.read_csv(
        root / "outputs" / "tables" / "stage4_cell_state_external_consistency.csv"
    )
    selected = [
        "hypoxia_tumor_cells",
        "proliferative_tumor_cells",
        "malignant_epithelial_cells",
        "emt_like_tumor_cells",
        "caf",
        "dendritic_cells",
        "b_cells",
        "plasma_cells",
    ]
    merged = correlations.merge(
        consistency[
            [
                "signature_name",
                "geo_cohorts_direction_match_tcga",
                "geo_cohorts_evaluable",
                "external_direction_consistency",
            ]
        ],
        on="signature_name",
        how="left",
    ).set_index("signature_name")
    cell_rows = []
    for name in selected:
        row = merged.loc[name]
        cell_rows.append(
            [
                name,
                f"{row['spearman_rho']:.3f}",
                _format_p(float(row["q_value_bh"])),
                row["direction"],
                f"{int(row['geo_cohorts_direction_match_tcga'])}/"
                f"{int(row['geo_cohorts_evaluable'])}",
                row["external_direction_consistency"],
            ]
        )
    _add_table(
        document,
        "Table 3. Selected cell-state associations with the frozen RNA risk score.",
        ["Signature", "TCGA rho", "BH q", "Direction", "GEO matches", "Consistency"],
        cell_rows,
    )

    _add_table(
        document,
        "Table 4. Orthogonal protein evidence and interpretation limits.",
        ["Evidence layer", "Samples/events", "Coverage", "Principal result", "Boundary"],
        [
            [
                "HPA",
                "Not a survival cohort",
                f"{evidence.hpa_supported_count}/{evidence.hpa_candidate_count} candidates",
                "Qualitative/IHC links available",
                "No structured quantitative IHC endpoint",
            ],
            [
                "CPTAC candidate proteins",
                f"{evidence.cptac_samples} tumors; {evidence.cptac_deaths} deaths",
                f"{evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total}",
                "LDHA, MKI67 and CDK1 direction-consistent nominal stage associations",
                "Exploratory; limited events",
            ],
            [
                "CPTAC hypoxia module",
                f"{evidence.cptac_usable_os} usable OS; {evidence.cptac_deaths} deaths",
                "Prespecified protein module",
                (
                    f"HR {evidence.cptac_hypoxia_hr:.2f}, 95% CI "
                    f"{evidence.cptac_hypoxia_ci_low:.2f}-"
                    f"{evidence.cptac_hypoxia_ci_high:.2f}; "
                    f"FDR={evidence.cptac_hypoxia_fdr:.4f}"
                ),
                "Exploratory quantitative support",
            ],
        ],
    )


def generate_main_manuscript(root: str | Path = ".") -> Path:
    paths = stage9_paths(root)
    paths["manuscript"].mkdir(parents=True, exist_ok=True)
    evidence = load_stage8_evidence(root)
    document = _new_document()
    _add_title(document)

    _add_heading(document, "Abstract", 1)
    for label, text in _abstract_text(evidence).items():
        _add_paragraph(
            document,
            f"{label}: {text}",
            bold_prefix=f"{label}:",
            first_line=False,
        )
    _add_paragraph(
        document,
        "Keywords: lung adenocarcinoma; overall survival; transcriptomics; "
        "external validation; cell-state signatures; proteomics",
        bold_prefix="Keywords:",
        first_line=False,
    )

    _add_heading(document, "Introduction", 1)
    _add_paragraph(
        document,
        "Lung adenocarcinoma (LUAD) is biologically heterogeneous, and patients "
        "with apparently similar clinicopathological characteristics can have "
        "substantially different outcomes. Stage, age and other routine variables "
        "remain essential prognostic factors, yet they incompletely represent the "
        "molecular programs that accompany aggressive disease. Transcriptome-wide "
        "profiling offers a route to quantify those programs, but the translational "
        "value of a molecular score depends on reproducibility beyond the cohort "
        "in which it was developed.",
    )
    _add_paragraph(
        document,
        "Many published prognostic signatures are vulnerable to optimistic "
        "estimation when feature selection or normalization precedes data "
        "splitting. Additional concerns include small external cohorts, "
        "platform-specific score distributions and the tendency to equate "
        "discrimination with calibrated clinical utility. A model can preserve "
        "risk ranking across cohorts while failing to transport an absolute "
        "decision threshold. We therefore treated leakage control, locked "
        "external evaluation and explicit reporting of threshold failure as "
        "central design requirements.",
    )
    _add_paragraph(
        document,
        "Interpretation presents a separate challenge. Bulk RNA risk scores mix "
        "malignant-cell states with immune and stromal signals. Curated signatures "
        "informed by single-cell atlases can characterize this mixture, but such "
        "scores remain inferences from bulk tissue unless raw single-cell profiles "
        "and linked outcomes are analysed directly. Orthogonal resources, "
        "including the Human Protein Atlas (HPA) and Clinical Proteomic Tumor "
        "Analysis Consortium (CPTAC), can add complementary evidence provided "
        "that qualitative immunohistochemistry links and limited-event proteomic "
        "analyses are not overstated.",
    )
    _add_paragraph(
        document,
        "Here, we developed a clinical plus RNA principal-component Cox model in "
        "TCGA-LUAD under nested cross-validation, froze the complete processing "
        "and prediction pipeline, and applied it without refitting to four GEO "
        "cohorts. We then related the frozen risk score to prespecified LUAD "
        "cell-state signatures and examined bounded HPA and CPTAC evidence. A "
        "digital pathology pilot was retained as an exploratory feasibility "
        "analysis because it did not show stable improvement over the clinical "
        "reference.",
    )

    _add_heading(document, "Methods", 1)
    _add_heading(document, "Study design and public data sources", 2)
    _add_paragraph(
        document,
        "This retrospective computational study used de-identified public data. "
        "TCGA-LUAD clinical and RNA-sequencing data were obtained through the "
        "Genomic Data Commons (GDC). External transcriptomic validation used "
        "GSE31210, GSE50081, GSE72094 and GSE68465 from the Gene Expression "
        "Omnibus (GEO). Biological interpretation used prespecified curated "
        "cell-state signatures, HPA gene and lung-cancer pathology pages, and the "
        "PDC000153 CPTAC LUAD Discovery Proteome. The prediction model was "
        "developed only in TCGA; GEO outcomes were not used to choose genes, "
        "components, coefficients or preprocessing parameters.",
    )
    _add_paragraph(
        document,
        "Evidence layers were kept analytically distinct. TCGA nested "
        "cross-validation and locked GEO evaluation addressed prognostic "
        "discrimination. Cell-state scoring characterized associations with the "
        "risk score. HPA and CPTAC provided orthogonal protein-level context. The "
        "WSI pilot addressed engineering feasibility and was not incorporated "
        "into the frozen prognostic model.",
    )

    _add_heading(document, "TCGA-LUAD cohort and survival endpoints", 2)
    _add_paragraph(
        document,
        f"Patient identifiers were standardized to the first 12 characters of the "
        f"TCGA barcode. The analysis included {evidence.tcga_patients} patients "
        f"with primary-tumor RNA-seq and usable overall survival (OS), including "
        f"{evidence.tcga_deaths} deaths and {evidence.tcga_censored} censored "
        "observations. OS time was represented in days, and status was encoded as "
        "1 for death and 0 for censoring. Records with missing or non-positive OS "
        "time or missing status were excluded. OS was the sole model-development "
        "endpoint.",
    )

    _add_heading(document, "RNA-seq preprocessing and clinical variables", 2)
    _add_paragraph(
        document,
        "GDC STAR-count files were restricted to primary tumors. TPM values were "
        "taken from the tpm_unstranded field, duplicate primary-tumor files were "
        "resolved deterministically, and protein-coding genes were retained. "
        "Expression was transformed as log2(TPM + 1). The source matrix contained "
        "19,962 protein-coding genes across 517 unique patients before survival "
        "matching. Missingness filtering, variance ranking, standardization and "
        "dimensionality reduction were fitted within training data only.",
    )
    _add_paragraph(
        document,
        "The clinical component used the harmonized covariates available in the "
        "locked Stage 2 pipeline, including age, sex and pathological stage where "
        "available. Categorical variables were encoded from training data, and "
        "missing-value handling was learned without access to validation or test "
        "folds. The final RNA representation comprised 25 principal components "
        "derived from 1,000 training-selected high-variance genes.",
    )

    _add_heading(
        document, "Prognostic model development and internal validation", 2
    )
    _add_paragraph(
        document,
        "Candidate clinical, RNA-only and combined models were compared using "
        "nested cross-validation. The outer loop used five folds and the inner "
        "loop used three folds, repeated for random seeds 42, 3407 and 2026. "
        "Stratification balanced event status and survival-time strata where "
        "feasible. Every preprocessing operation, including gene filtering, "
        "scaling, PCA and hyperparameter selection, was fitted in the applicable "
        "training fold. Outer-fold predictions were retained as out-of-fold risk "
        "scores.",
    )
    _add_paragraph(
        document,
        "The prespecified frozen model was the RNA PCA_25 plus clinical Cox model. "
        "Clinical-only Cox was the principal reference. Elastic-net Cox and "
        "DeepSurv configurations were evaluated during robustness analyses; "
        "DeepSurv showed larger train-to-test gaps and was not selected. After "
        "model selection, training-gene identities, transformations, PCA loadings "
        "and Cox coefficients were fixed for external application.",
    )

    _add_heading(document, "External GEO validation", 2)
    _add_paragraph(
        document,
        "Each GEO cohort was imported and normalized independently. Probe "
        "identifiers were mapped to gene symbols and duplicate probes were "
        "collapsed using the prespecified pipeline. Expression standardization "
        "was performed within each external cohort to address platform scale, but "
        "GEO data were never pooled with TCGA. Missing frozen-model genes were "
        "imputed according to the locked external-validation procedure, and the "
        "fixed TCGA feature order, PCA transformation and Cox coefficients were "
        "then applied without outcome-guided modification.",
    )
    _add_paragraph(
        document,
        "Primary external evaluation used the continuous risk score. We estimated "
        "Harrell's C-index, time-dependent AUC at 1, 3 and 5 years, and the hazard "
        "ratio per standard deviation of risk score. The TCGA-derived cutoff was "
        "tested as a transportability analysis. Cohort-median Kaplan-Meier "
        "comparisons were treated as sensitivity analyses because they use an "
        "external-cohort-specific threshold.",
    )

    _add_heading(document, "Cell-state signature scoring", 2)
    _add_paragraph(
        document,
        f"We evaluated {evidence.signature_count} prespecified signatures covering "
        "malignant epithelial, EMT-like, proliferative and hypoxic tumor states; "
        "T-cell, macrophage, dendritic, B-cell, plasma-cell, NK-cell and mast-cell "
        "contexts; cancer-associated fibroblasts; and endothelial cells. "
        "Signatures were curated before outcome association analysis and were not "
        "optimized against survival. For each dataset, gene expression was "
        "standardized within that dataset and signature scores were calculated as "
        "the mean z score of available member genes.",
    )
    _add_paragraph(
        document,
        "In TCGA, Spearman correlations quantified association between each "
        "cell-state score and the frozen RNA risk score. Median-defined high- and "
        "low-risk groups were compared using nonparametric tests. Exploratory "
        "univariable and age-, sex- and stage-adjusted Cox models assessed "
        "prognostic associations. In GEO, analyses were limited to risk-score "
        "correlations and directional consistency. These analyses infer "
        "cell-state context from bulk expression and do not replace survival "
        "analysis in a raw single-cell cohort.",
    )

    _add_heading(document, "HPA protein/IHC evidence", 2)
    _add_paragraph(
        document,
        f"We prespecified {evidence.hpa_candidate_count} genes representing "
        "hypoxia, proliferation, EMT-like, CAF/matrix and dendritic/B/plasma "
        "programs. HPA records were queried for protein evidence, antibodies and "
        "lung-cancer pathology links. Because image-level lung-cancer staining "
        "was not available as a uniform structured quantitative endpoint, HPA "
        "findings were classified as qualitative or IHC-link evidence and were "
        "not analysed as an independent survival-validation cohort.",
    )

    _add_heading(document, "CPTAC-LUAD quantitative proteomic analysis", 2)
    _add_paragraph(
        document,
        f"The PDC000153 discovery proteome comprised {evidence.cptac_samples} "
        f"primary tumors and {evidence.cptac_proteins:,} protein genes. Duplicate "
        "aliquots were collapsed by patient-level mean where required. Protein "
        "abundance values were standardized for candidate and module analyses. "
        f"Usable OS was available for {evidence.cptac_usable_os} patients, with "
        f"{evidence.cptac_deaths} deaths. Of the prespecified candidates, "
        f"{evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} "
        "were quantified.",
    )
    _add_paragraph(
        document,
        "Candidate proteins were assessed for direction-consistent association "
        "with clinical stage and for exploratory survival association. "
        "Prespecified mechanism-layer protein scores were calculated from "
        "available proteins and tested using Cox models. Benjamini-Hochberg false "
        "discovery rates were reported. The low event count and incomplete "
        "candidate coverage make these analyses exploratory quantitative "
        "proteomic support rather than conclusive validation.",
    )

    _add_heading(document, "Exploratory WSI feasibility analysis", 2)
    _add_paragraph(
        document,
        "A separate 100-slide TCGA-LUAD pilot evaluated OpenSlide reading, tissue "
        "segmentation, 256-pixel patch extraction, pretrained ResNet50 feature "
        "extraction and attention-based multiple-instance learning. Clinical-only, "
        "pathology-only and clinical-pathology models were compared on a small "
        "held-out pilot split. No WSI variables entered the frozen RNA-clinical "
        "model, and the pilot was not expanded after pathology models failed to "
        "show stable improvement.",
    )

    _add_heading(document, "Statistical analysis", 2)
    _add_paragraph(
        document,
        "Discrimination was summarized by C-index and time-dependent AUC. "
        "Kaplan-Meier curves and log-rank tests described group separation, while "
        "Cox models reported hazard ratios with 95% confidence intervals. "
        "Spearman coefficients quantified risk-signature association. Multiple "
        "testing used Benjamini-Hochberg adjustment where families of signatures "
        "or proteins were evaluated. All tests were two-sided. Analyses were "
        "implemented in Python using pandas, NumPy, scikit-learn, lifelines, "
        "PyTorch and project-specific modules.",
    )
    _add_paragraph(
        document,
        "Detailed feature-space comparisons are provided in Additional file 1. "
        "The locked-model specification and external input coverage are provided "
        "in Additional file 2. Cell-state definitions and complete analyses are "
        "provided in Additional file 3. HPA qualitative/IHC-link evidence is "
        "provided in Additional file 4. CPTAC candidate and module analyses are "
        "provided in Additional file 5. External calibration and cutoff "
        "diagnostics are provided in Additional file 6. The WSI feasibility "
        "pilot is provided in Additional file 7.",
    )

    _add_heading(document, "Results", 1)
    _add_heading(document, "Cohort construction and model development", 2)
    _add_paragraph(
        document,
        f"The study workflow separated TCGA development, locked GEO validation "
        f"and post hoc biological interpretation (Fig. 1). The TCGA cohort "
        f"included {evidence.tcga_patients} patients, "
        f"{evidence.tcga_deaths} deaths and {evidence.tcga_censored} censored "
        "observations; external cohorts contributed "
        f"{evidence.external_total} patients (Table 1).",
    )
    _add_paragraph(
        document,
        "Across 15 outer folds, the RNA PCA_25 plus clinical Cox model achieved "
        f"a mean C-index of {evidence.nested_combined_cindex:.3f} "
        f"(SD {evidence.nested_combined_sd:.3f}), compared with "
        f"{evidence.nested_clinical_cindex:.3f} "
        f"(SD {evidence.nested_clinical_sd:.3f}) for clinical Cox. Mean 1-, 3- "
        f"and 5-year AUCs were {evidence.nested_auc_1y:.3f}, "
        f"{evidence.nested_auc_3y:.3f} and {evidence.nested_auc_5y:.3f}, "
        "respectively. The absolute C-index difference relative to the clinical "
        "reference was small. DeepSurv configurations had larger train-to-test "
        "gaps and did not support a neural-network advantage (Table 2; Fig. 2).",
    )

    _add_heading(document, "External validation across four GEO cohorts", 2)
    external_sentences = []
    for row in _external_rows(evidence):
        external_sentences.append(
            f"{row['cohort']} (n={row['n']}, {row['deaths']} deaths) "
            f"C-index {row['c_index']:.3f}, HR per SD {row['hr']:.3f}, "
            f"P={_format_p(row['p'])}"
        )
    _add_paragraph(
        document,
        "The complete frozen model was applied without refitting. External "
        "results were: "
        + "; ".join(external_sentences)
        + ". The continuous risk score was therefore associated with OS in all "
        "four cohorts despite differences in platform and gene coverage.",
    )
    _add_paragraph(
        document,
        "The TCGA-derived median cutoff did not reliably separate survival in "
        "external cohorts and assigned all samples to one group in two cohorts. "
        "Cohort-median sensitivity analyses produced clearer Kaplan-Meier "
        "separation but do not establish a transferable threshold. The external "
        "evidence consequently supports relative risk ranking rather than "
        "calibrated absolute-risk classification (Fig. 3).",
    )

    _add_heading(
        document, "Cell-state programs associated with RNA risk score", 2
    )
    _add_paragraph(
        document,
        "In TCGA, the strongest positive correlations with risk were hypoxia "
        "(rho=0.354), proliferation (rho=0.349), malignant epithelial state "
        "(rho=0.277), EMT-like state (rho=0.193) and CAF (rho=0.104). Hypoxia, "
        "proliferation, EMT-like and CAF scores were higher in the high-risk "
        "group and showed concordant correlation directions across all four GEO "
        "cohorts. Dendritic-cell, B-cell and plasma-cell scores supported a "
        "lower-risk immune context (Table 3; Fig. 4).",
    )
    _add_paragraph(
        document,
        "Adjusted exploratory Cox analyses identified nominal associations for "
        "proliferative, hypoxic, dendritic, B-cell, plasma-cell and mast-cell "
        "scores. Findings for M2 macrophage, regulatory T-cell, exhausted CD8, "
        "cytotoxic CD8, M1 macrophage and NK signatures were mixed or externally "
        "unstable and were not included in the core interpretation. These results "
        "describe cell-state associations inferred from bulk RNA rather than "
        "direct measurement of cell frequencies or an established biological "
        "mechanism.",
    )

    _add_heading(
        document, "HPA and CPTAC protein-level orthogonal support", 2
    )
    _add_paragraph(
        document,
        f"All {evidence.hpa_supported_count} prespecified genes had an HPA record "
        "or lung-cancer pathology link, providing qualitative/IHC-link evidence. "
        "The lack of a consistent structured quantitative staining endpoint "
        "precluded formal effect-size estimation from HPA.",
    )
    _add_paragraph(
        document,
        f"In CPTAC, {evidence.cptac_candidates_matched}/"
        f"{evidence.cptac_candidate_total} candidates were quantified. LDHA, "
        "MKI67 and CDK1 showed nominally significant stage associations in the "
        "prespecified direction. No individual candidate protein retained a "
        "survival association after FDR correction. The hypoxia protein module "
        f"was associated with poorer OS (HR {evidence.cptac_hypoxia_hr:.3f} per "
        f"SD, 95% CI {evidence.cptac_hypoxia_ci_low:.3f}-"
        f"{evidence.cptac_hypoxia_ci_high:.3f}, P="
        f"{_format_p(evidence.cptac_hypoxia_p)}, FDR="
        f"{evidence.cptac_hypoxia_fdr:.4f}). Given only "
        f"{evidence.cptac_deaths} deaths, these data provide exploratory "
        "quantitative proteomic support, particularly for the hypoxia and "
        "proliferation axes (Table 4; Fig. 5).",
    )

    _add_heading(document, "Exploratory WSI feasibility findings", 2)
    _add_paragraph(
        document,
        f"The 100-slide pilot completed patch extraction, GPU feature extraction "
        f"and survival-model training. The held-out clinical-only C-index was "
        f"{evidence.wsi_clinical_cindex:.3f}, compared with "
        f"{evidence.wsi_pathology_cindex:.3f} for the best pathology-only model "
        f"and {evidence.wsi_fusion_cindex:.3f} for the best fusion model. These "
        "small-pilot estimates are not comparative efficacy estimates. They show "
        "technical feasibility while providing no stable evidence that WSI added "
        "predictive value (Additional file 7).",
    )

    _add_heading(document, "Discussion", 1)
    _add_paragraph(
        document,
        "The central finding is that a frozen TCGA-derived transcriptomic-clinical "
        "score retained prognostic ranking across four independent GEO cohorts. "
        "This reproducibility was obtained without refitting the selected genes, "
        "PCA representation or Cox coefficients to external outcomes. The "
        "external C-indices ranged from 0.624 to 0.700, and the continuous score "
        "was associated with OS in every cohort.",
    )
    _add_paragraph(
        document,
        "The result should nevertheless be interpreted against the clinical "
        "reference. In TCGA nested cross-validation, adding the RNA representation "
        "increased mean C-index from 0.629 to 0.636. This modest difference argues "
        "against portraying transcriptomics as a large replacement for standard "
        "clinical information. Instead, the molecular component may be most useful "
        "as a reproducible research stratifier whose value must be tested under "
        "prospective assay and treatment conditions.",
    )
    _add_paragraph(
        document,
        "The failure of the TCGA cutoff to transfer is also informative. Gene "
        "coverage, platform scaling, patient composition and treatment context "
        "differed among GEO datasets. Cohort-specific ranking was substantially "
        "more stable than absolute score location. Future work should therefore "
        "include assay harmonization, prospective calibration and prespecified "
        "decision thresholds rather than importing a retrospective median cutoff.",
    )
    _add_paragraph(
        document,
        "The cell-state analysis connected higher risk to hypoxic, proliferative, "
        "EMT-like and fibroblast-associated programs, with lower-risk directionality "
        "for dendritic, B-cell and plasma-cell contexts. Concordance across GEO "
        "cohorts strengthens the interpretation that the score tracks coherent "
        "tumor and microenvironment programs. However, curated signatures cannot "
        "resolve cell abundance from cell-intrinsic expression, and the present "
        "study did not analyse a raw single-cell cohort with linked survival.",
    )
    _add_paragraph(
        document,
        "Protein resources added a bounded orthogonal layer. HPA established that "
        "the candidate proteins and lung-cancer pathology links can be inspected, "
        "but its contribution here was qualitative. CPTAC quantified most "
        "candidates and prioritized LDHA, MKI67 and CDK1 by direction-consistent "
        "stage associations. The hypoxia module showed the clearest survival "
        "signal. These observations are compatible with the RNA-based interpretation "
        "but remain exploratory because the survival analysis contained only 26 "
        "deaths.",
    )
    _add_paragraph(
        document,
        "The neutral WSI pilot further defined the scope of the study. Although "
        "the pathology pipeline was operational, neither pathology-only nor fusion "
        "models produced stable improvement over the clinical comparator. This "
        "result supports keeping digital pathology outside the principal claim and "
        "prioritizing larger, rigorously powered pathology studies if that direction "
        "is revisited.",
    )
    _add_paragraph(
        document,
        "Together, these findings position the model as an externally reproducible "
        "research score with biologically interpretable associations. Translation "
        "will require prospective validation, standardized expression assays, "
        "calibration, predefined thresholds and demonstration that the score adds "
        "decision-relevant information beyond routine clinicopathological factors.",
    )

    _add_heading(document, "Limitations", 1)
    limitations = [
        "The study was retrospective and relied on public cohorts with heterogeneous treatment, follow-up and clinical annotation.",
        "The RNA component added only a modest mean C-index increment over clinical-only Cox in TCGA nested cross-validation.",
        "External validation used microarray cohorts with incomplete frozen-gene coverage and platform-dependent score distributions.",
        "The TCGA-derived cutoff was not portable, and prospective calibration or decision-impact evaluation was not performed.",
        "Cell states were inferred using curated signatures from bulk RNA rather than measured in raw single-cell data linked to survival.",
        "HPA evidence was qualitative and link based, without a harmonized quantitative lung-cancer IHC endpoint.",
        f"CPTAC survival analyses included only {evidence.cptac_deaths} deaths and quantified {evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} candidates.",
        "The WSI experiment was a small feasibility pilot and did not show stable added predictive value.",
        "The study did not test treatment interaction, prospective utility or implementation in a clinical workflow.",
    ]
    for item in limitations:
        _add_paragraph(
            document, item, style="List Number", first_line=False
        )

    _add_heading(document, "Conclusions", 1)
    _add_paragraph(
        document,
        "A leakage-controlled TCGA transcriptomic-clinical Cox model reproduced "
        "survival risk ranking across four independent GEO cohorts. Higher risk "
        "was associated with hypoxia, proliferation, EMT-like tumor and CAF "
        "programs, while HPA and exploratory CPTAC analyses offered bounded "
        "protein-level context. The limited RNA increment over clinical variables, "
        "non-portable cutoff, signature-based cell-state inference and low CPTAC "
        "event count preclude claims of clinical readiness or proof of mechanism. "
        "Prospective and assay-harmonized studies are required.",
    )

    _add_heading(document, "List of abbreviations", 1)
    abbreviations = [
        ("AUC", "area under the time-dependent receiver operating characteristic curve"),
        ("CAF", "cancer-associated fibroblast"),
        ("C-index", "concordance index"),
        ("CPTAC", "Clinical Proteomic Tumor Analysis Consortium"),
        ("FDR", "false discovery rate"),
        ("GDC", "Genomic Data Commons"),
        ("GEO", "Gene Expression Omnibus"),
        ("HPA", "Human Protein Atlas"),
        ("IHC", "immunohistochemistry"),
        ("LUAD", "lung adenocarcinoma"),
        ("OS", "overall survival"),
        ("PCA", "principal component analysis"),
        ("PDC", "Proteomic Data Commons"),
        ("TCGA", "The Cancer Genome Atlas"),
        ("TPM", "transcripts per million"),
        ("WSI", "whole-slide image"),
    ]
    for short, long in abbreviations:
        _add_paragraph(
            document, f"{short}: {long}", bold_prefix=f"{short}:", first_line=False
        )

    _add_heading(document, "Declarations", 1)
    declaration_sections = declaration_text()
    for heading, text in declaration_sections.items():
        _add_heading(document, heading, 2)
        _add_paragraph(document, text)

    _add_heading(document, "References", 1)
    references = [
        "1. The Cancer Genome Atlas Research Network. Comprehensive molecular profiling of lung adenocarcinoma. Nature. 2014;511:543-550. doi:10.1038/nature13385.",
        "2. Cox DR. Regression models and life-tables. J R Stat Soc Series B Stat Methodol. 1972;34:187-220.",
        "3. Harrell FE Jr, Lee KL, Mark DB. Multivariable prognostic models: issues in developing models, evaluating assumptions and adequacy, and measuring and reducing errors. Stat Med. 1996;15:361-387.",
        "4. Heagerty PJ, Lumley T, Pepe MS. Time-dependent ROC curves for censored survival data and a diagnostic marker. Biometrics. 2000;56:337-344.",
        "5. Edgar R, Domrachev M, Lash AE. Gene Expression Omnibus: NCBI gene expression and hybridization array data repository. Nucleic Acids Res. 2002;30:207-210. doi:10.1093/nar/30.1.207.",
        "6. Shedden K, Taylor JMG, Enkemann SA, et al. Gene expression-based survival prediction in lung adenocarcinoma: a multi-site, blinded validation study. Nat Med. 2008;14:822-827. doi:10.1038/nm.1790.",
        "7. Uhlen M, Zhang C, Lee S, et al. A pathology atlas of the human cancer transcriptome. Science. 2017;357:eaan2507. doi:10.1126/science.aan2507.",
        "8. Gillette MA, Satpathy S, Cao S, et al. Proteogenomic characterization reveals therapeutic vulnerabilities in lung adenocarcinoma. Cell. 2020;182:200-225.e35. doi:10.1016/j.cell.2020.06.013.",
        "9. Subramanian A, Tamayo P, Mootha VK, et al. Gene set enrichment analysis: a knowledge-based approach for interpreting genome-wide expression profiles. Proc Natl Acad Sci USA. 2005;102:15545-15550. doi:10.1073/pnas.0506580102.",
        "10. [AUTHOR TO COMPLETE: add the primary LUAD single-cell atlas references used to justify the curated cell-state marker sets.]",
        "11. [AUTHOR TO COMPLETE: add the original cohort publications for GSE31210, GSE50081 and GSE72094 after bibliographic verification.]",
    ]
    for reference in references:
        _add_paragraph(document, reference, first_line=False)

    _add_heading(document, "Tables", 1)
    _build_main_tables(document, evidence, paths["root"])

    _add_heading(document, "Figure legends", 1)
    figure_legends = [
        (
            "Figure 1. Study design and evidence workflow. "
            "TCGA-LUAD model development and nested cross-validation were "
            "separated from locked GEO validation. Cell-state, HPA and CPTAC "
            "analyses were interpretation layers; WSI was exploratory."
        ),
        (
            "Figure 2. Internal model comparison. Nested-cross-validation "
            "C-indices and feature-space results show a modest increment for RNA "
            "PCA_25 plus clinical Cox over clinical Cox and larger overfitting "
            "for complex neural models."
        ),
        (
            "Figure 3. External validation. The frozen continuous risk score "
            "retained discrimination and survival association across four GEO "
            "cohorts, while the TCGA-derived absolute cutoff did not transfer "
            "reliably."
        ),
        (
            "Figure 4. Cell-state interpretation. TCGA correlations, risk-group "
            "comparisons and GEO direction checks associate higher risk with "
            "hypoxia, proliferation, EMT-like and CAF programs and lower risk "
            "with dendritic/B/plasma context."
        ),
        (
            "Figure 5. Orthogonal protein evidence. HPA provides qualitative/IHC "
            "links; exploratory CPTAC analyses summarize candidate coverage, "
            "direction-consistent stage associations and the hypoxia module "
            "survival result."
        ),
    ]
    for legend in figure_legends:
        _add_paragraph(document, legend, first_line=False)

    document.save(paths["main"])
    return paths["main"]


def declaration_text() -> dict[str, str]:
    return {
        "Ethics approval and consent to participate": (
            "This study used de-identified data obtained from public research "
            "repositories and recruited no new participants. [AUTHOR TO CONFIRM: "
            "insert the institutional ethics determination, committee name and "
            "approval/waiver reference if required by the authors' institution.]"
        ),
        "Consent for publication": (
            "Not applicable; this manuscript contains no identifiable individual "
            "participant data. [AUTHOR TO CONFIRM.]"
        ),
        "Availability of data and materials": (
            "TCGA-LUAD data are available through the NCI Genomic Data Commons "
            "(https://portal.gdc.cancer.gov/projects/TCGA-LUAD). GEO validation "
            "data are available under GSE31210, GSE50081, GSE72094 and GSE68465 "
            "(https://www.ncbi.nlm.nih.gov/geo/). HPA evidence is available from "
            "https://www.proteinatlas.org/. CPTAC-LUAD proteomics are available "
            "from the Proteomic Data Commons under PDC000153 "
            "(https://pdc.cancer.gov/). [AUTHOR TO COMPLETE: provide the public "
            "code repository URL, archived release DOI/version and the location "
            "of derived analysis tables permitted for redistribution.]"
        ),
        "Competing interests": (
            "[AUTHOR TO CONFIRM: declare that the authors have no competing "
            "interests, or provide the complete disclosure.]"
        ),
        "Funding": (
            "[AUTHOR TO COMPLETE: list funders, grant numbers and each funder's "
            "role in study design, data analysis, manuscript preparation and the "
            "decision to submit. If there was no specific funding, state this "
            "explicitly after author confirmation.]"
        ),
        "Authors' contributions": (
            "[AUTHOR TO COMPLETE using CRediT-compatible initials: "
            "conceptualization; methodology; software; validation; formal "
            "analysis; data curation; visualization; writing-original draft; "
            "writing-review and editing; supervision; funding acquisition. All "
            "authors must approve the submitted version and accept accountability.]"
        ),
        "Acknowledgements": (
            "The authors acknowledge the participants and investigators who "
            "generated and shared TCGA, GEO, HPA, CPTAC and PDC resources. "
            "[AUTHOR TO COMPLETE: add non-author contributions and required "
            "acknowledgement wording.]"
        ),
        "Authors' information": (
            "[AUTHOR TO COMPLETE: provide relevant author information only if "
            "needed, including ORCID identifiers.]"
        ),
    }


def generate_declarations(root: str | Path = ".") -> Path:
    paths = stage9_paths(root)
    paths["manuscript"].mkdir(parents=True, exist_ok=True)
    document = _new_document()
    _add_title(document, "Declarations")
    _add_paragraph(document, f"Manuscript title: {TITLE}", first_line=False)
    for heading, text in declaration_text().items():
        _add_heading(document, heading, 1)
        _add_paragraph(document, text)
    document.save(paths["declarations"])
    return paths["declarations"]


def _main_word_count(path: Path) -> int:
    if not path.exists():
        return 0
    from docx import Document

    document = Document(path)
    text = " ".join(paragraph.text for paragraph in document.paragraphs)
    return len(re.findall(r"\b[\w'-]+\b", text))


def generate_title_page(root: str | Path = ".") -> Path:
    paths = stage9_paths(root)
    paths["manuscript"].mkdir(parents=True, exist_ok=True)
    document = _new_document()
    _add_title(document)
    _add_paragraph(
        document, f"Running title: {RUNNING_TITLE}", first_line=False
    )
    _add_heading(document, "Authors", 1)
    _add_paragraph(
        document,
        "[AUTHOR TO COMPLETE: Full Name 1]1, [AUTHOR TO COMPLETE: Full Name 2]2, "
        "[AUTHOR TO COMPLETE: Full Name 3]1,*",
        first_line=False,
    )
    _add_heading(document, "Affiliations", 1)
    _add_paragraph(
        document,
        "1 [AUTHOR TO COMPLETE: Department, Institution, City, Postal code, Country]",
        first_line=False,
    )
    _add_paragraph(
        document,
        "2 [AUTHOR TO COMPLETE: Department, Institution, City, Postal code, Country]",
        first_line=False,
    )
    _add_heading(document, "Corresponding author", 1)
    _add_paragraph(
        document,
        "[AUTHOR TO COMPLETE: Name, full postal address, telephone if requested, "
        "email and ORCID]",
        first_line=False,
    )
    _add_heading(document, "Manuscript information", 1)
    _add_paragraph(
        document,
        f"Article type: Research Article\n"
        f"Target journal: Journal of Translational Medicine\n"
        f"Main-manuscript word count: {_main_word_count(paths['main'])}\n"
        "Abstract word count: [generated compliance report]\n"
        "Main tables: 4\n"
        "Main figures: 5\n"
        "Additional files cited: 7",
        first_line=False,
    )
    _add_heading(document, "Author metadata still required", 1)
    for item in [
        "Final author order and spelling",
        "Affiliation mapping",
        "Corresponding-author details",
        "ORCID identifiers",
        "Equal-contribution or consortium statements, if applicable",
    ]:
        _add_paragraph(document, item, style="List Bullet", first_line=False)
    document.save(paths["title_page"])
    return paths["title_page"]


def generate_cover_letter(root: str | Path = ".") -> Path:
    paths = stage9_paths(root)
    paths["manuscript"].mkdir(parents=True, exist_ok=True)
    evidence = load_stage8_evidence(root)
    document = _new_document(letter=True)
    _add_paragraph(
        document,
        f"{date.today().strftime('%d %B %Y')}",
        first_line=False,
    )
    _add_paragraph(
        document,
        "Editor-in-Chief\nJournal of Translational Medicine",
        first_line=False,
    )
    _add_paragraph(document, "Dear Editor,", first_line=False)
    _add_paragraph(
        document,
        f"We are pleased to submit the Research Article entitled \"{TITLE}\" for "
        "consideration in Journal of Translational Medicine.",
    )
    _add_paragraph(
        document,
        "The manuscript addresses a translationally relevant problem: whether a "
        "molecular survival score can remain reproducible across independent lung "
        "adenocarcinoma cohorts while retaining a biologically interpretable and "
        "appropriately bounded evidence chain. The study combines leakage-controlled "
        "TCGA development with locked validation in four GEO cohorts, followed by "
        "curated cell-state interpretation and orthogonal HPA and CPTAC analyses. "
        "This design aligns with the journal's interest in work that connects "
        "molecular observations with clinically relevant disease stratification.",
    )
    _add_paragraph(
        document,
        f"The frozen score was evaluated in {evidence.external_total} external "
        "patients and showed C-indices of 0.624-0.700. Its novelty lies less in "
        "algorithmic complexity than in the combination of multi-cohort locked "
        "validation, explicit assessment of cutoff transportability, cell-state "
        "associations inferred from bulk RNA, qualitative HPA/IHC links and "
        "exploratory quantitative CPTAC support. We report the modest increment "
        "over clinical-only Cox and do not present the model as suitable for "
        "current clinical decision-making.",
    )
    _add_paragraph(
        document,
        f"The CPTAC analysis included only {evidence.cptac_deaths} deaths and is "
        "therefore described as exploratory. Likewise, the WSI pilot established "
        "technical feasibility but did not show stable improvement over the "
        "clinical comparator; it is included as a neutral supplementary finding "
        "rather than as a principal claim.",
    )
    _add_paragraph(
        document,
        "We confirm that this manuscript is not under consideration by another "
        "journal and has not been published previously, except for any disclosed "
        "preprint [AUTHOR TO CONFIRM OR DELETE]. All authors have reviewed and "
        "approved the submitted manuscript and agree to be accountable for the "
        "work [AUTHOR TO CONFIRM BEFORE SUBMISSION].",
    )
    _add_paragraph(
        document,
        "Competing interests: [AUTHOR TO CONFIRM: declare none or provide the "
        "complete disclosure]. Funding and author-contribution statements are "
        "provided in the declarations file and require final author verification.",
    )
    _add_paragraph(
        document,
        "Thank you for considering our manuscript. We believe that its external "
        "validation, transparent negative findings and restrained biological "
        "interpretation will be of interest to readers working at the interface "
        "of cancer genomics, biomarkers and translational oncology.",
    )
    _add_paragraph(document, "Sincerely,", first_line=False)
    _add_paragraph(
        document,
        "[AUTHOR TO COMPLETE: Corresponding author name]\n"
        "[AUTHOR TO COMPLETE: Institution]\n"
        "[AUTHOR TO COMPLETE: Email]",
        first_line=False,
    )
    document.save(paths["cover_letter"])
    return paths["cover_letter"]


def _document_text(path: Path) -> str:
    if not path.exists():
        return ""
    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _document_xml(path: Path) -> str:
    if not path.exists():
        return ""
    with ZipFile(path) as archive:
        names = [
            name
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]
        return "\n".join(
            archive.read(name).decode("utf-8", errors="ignore") for name in names
        )


def _citation_order(text: str, prefix: str, maximum: int) -> tuple[list[int], bool]:
    if prefix == "Fig":
        numbers = [int(value) for value in re.findall(r"\bFig\.\s*(\d+)", text)]
    elif prefix == "Table":
        numbers = [int(value) for value in re.findall(r"\bTable\s+(\d+)", text)]
    else:
        numbers = [
            int(value)
            for value in re.findall(r"\bAdditional file\s+(\d+)", text)
        ]
    first_seen: list[int] = []
    for number in numbers:
        if number not in first_seen:
            first_seen.append(number)
    return first_seen, first_seen[:maximum] == list(range(1, maximum + 1))


def build_submission_checklist(root: str | Path = ".") -> pd.DataFrame:
    paths = stage9_paths(root)
    rows = [
        {
            "category": "Required manuscript file",
            "submission_item": "Main manuscript",
            "expected_file": "outputs/manuscript/JTM_main_manuscript_draft.docx",
            "required": True,
            "generated": paths["main"].exists(),
            "status": "draft_generated" if paths["main"].exists() else "missing",
            "action_before_submission": (
                "Resolve author placeholders, verify bibliography, and insert final "
                "figure/table files."
            ),
        },
        {
            "category": "Required manuscript file",
            "submission_item": "Title page",
            "expected_file": "outputs/manuscript/JTM_title_page.docx",
            "required": True,
            "generated": paths["title_page"].exists(),
            "status": "metadata_pending",
            "action_before_submission": "Complete authors, affiliations and correspondence.",
        },
        {
            "category": "Editorial file",
            "submission_item": "Cover letter",
            "expected_file": "outputs/manuscript/JTM_cover_letter.docx",
            "required": True,
            "generated": paths["cover_letter"].exists(),
            "status": "author_confirmation_pending",
            "action_before_submission": "Confirm exclusivity, approval, preprint and conflicts.",
        },
        {
            "category": "Required statements",
            "submission_item": "Declarations",
            "expected_file": "outputs/manuscript/JTM_declarations.docx",
            "required": True,
            "generated": paths["declarations"].exists(),
            "status": "author_confirmation_pending",
            "action_before_submission": "Complete ethics, funding, contributions and conflicts.",
        },
    ]
    for number in range(1, 6):
        rows.append(
            {
                "category": "Main figure",
                "submission_item": f"Figure {number}",
                "expected_file": f"[AUTHOR TO ASSEMBLE: JTM_Figure_{number}.tif or .pdf]",
                "required": True,
                "generated": False,
                "status": "assembly_pending",
                "action_before_submission": (
                    "Assemble source panels, verify resolution, labels and legend."
                ),
            }
        )
    for number in range(1, 5):
        rows.append(
            {
                "category": "Main table",
                "submission_item": f"Table {number}",
                "expected_file": "Embedded as editable table in main manuscript",
                "required": True,
                "generated": paths["main"].exists(),
                "status": "draft_embedded",
                "action_before_submission": "Editorial and numerical proofread.",
            }
        )
    additional_labels = {
        1: "Nested-CV and feature-space results",
        2: "Frozen-model specification and external input coverage",
        3: "Cell-state signature definitions and complete analyses",
        4: "HPA qualitative/IHC-link evidence",
        5: "CPTAC candidate and module analyses",
        6: "External calibration and cutoff diagnostics",
        7: "Exploratory WSI feasibility analysis",
    }
    for number, label in additional_labels.items():
        rows.append(
            {
                "category": "Additional file",
                "submission_item": f"Additional file {number}: {label}",
                "expected_file": f"[AUTHOR TO ASSEMBLE: Additional_file_{number}]",
                "required": True,
                "generated": False,
                "status": "source_outputs_available_assembly_pending",
                "action_before_submission": (
                    "Assemble only de-identified, redistribution-permitted tables "
                    "and figures with a content legend."
                ),
            }
        )
    rows.extend(
        [
            {
                "category": "Reporting",
                "submission_item": "Reporting guideline checklist",
                "expected_file": "[AUTHOR TO SELECT: REMARK/TRIPOD-related checklist as applicable]",
                "required": False,
                "generated": False,
                "status": "editorial_decision_pending",
                "action_before_submission": "Select the closest applicable checklist.",
            },
            {
                "category": "Reproducibility",
                "submission_item": "Public code release",
                "expected_file": "[AUTHOR TO COMPLETE: repository URL and archived DOI]",
                "required": False,
                "generated": False,
                "status": "public_release_pending",
                "action_before_submission": "Create a versioned public release without raw large data.",
            },
            {
                "category": "Permissions",
                "submission_item": "Third-party figure permissions",
                "expected_file": "[AUTHOR TO CONFIRM]",
                "required": False,
                "generated": False,
                "status": "review_pending",
                "action_before_submission": (
                    "Use newly assembled plots or obtain permission for any reused content."
                ),
            },
        ]
    )
    frame = pd.DataFrame(rows)
    paths["tables"].mkdir(parents=True, exist_ok=True)
    frame.to_csv(paths["checklist"], index=False)
    return frame


def compliance_report(root: str | Path = ".") -> str:
    paths = stage9_paths(root)
    text = _document_text(paths["main"])
    xml = _document_xml(paths["main"])
    abstract_match = re.search(
        r"Abstract\s+(.*?)\s+Keywords:", text, flags=re.DOTALL | re.IGNORECASE
    )
    abstract = abstract_match.group(1) if abstract_match else ""
    abstract_words = len(re.findall(r"\b[\w'-]+\b", abstract))
    keyword_match = re.search(
        r"Keywords:\s*(.+)", text, flags=re.IGNORECASE
    )
    keywords = (
        [item.strip() for item in keyword_match.group(1).split(";") if item.strip()]
        if keyword_match
        else []
    )
    figures, figure_order = _citation_order(text, "Fig", 5)
    tables, table_order = _citation_order(text, "Table", 4)
    supplements, supplement_order = _citation_order(text, "Additional", 7)
    required_sections = [
        "Abstract",
        "Introduction",
        "Methods",
        "Results",
        "Discussion",
        "Limitations",
        "Conclusions",
        "List of abbreviations",
        "Declarations",
        "References",
    ]
    missing_sections = [section for section in required_sections if section not in text]
    placeholders = len(re.findall(r"\[AUTHOR TO (?:COMPLETE|CONFIRM)", text))
    checklist = build_submission_checklist(root)
    missing_required_files = checklist.loc[
        checklist["required"].astype(bool) & ~checklist["generated"].astype(bool),
        "submission_item",
    ].tolist()
    checks = [
        ("Structured abstract <=350 words", abstract_words <= 350 and abstract_words > 0),
        ("Three to ten keywords", 3 <= len(keywords) <= 10),
        ("Required manuscript sections present", not missing_sections),
        ("Continuous line numbering encoded", 'w:restart="continuous"' in xml),
        ("Page-number field encoded", " PAGE " in xml),
        ("Double line spacing encoded", 'w:line="480"' in xml or 'w:lineRule="auto"' in xml),
        ("Figures first cited in order 1-5", figure_order),
        ("Tables first cited in order 1-4", table_order),
        ("Additional files first cited in order 1-7", supplement_order),
    ]
    check_lines = "\n".join(
        f"| {name} | {'PASS' if passed else 'REVIEW'} |" for name, passed in checks
    )
    return f"""# Journal of Translational Medicine Compliance Report

Generated: {datetime.now().isoformat(timespec="seconds")}

## Scope

This audit compares the generated draft with the official Journal of Translational Medicine Research Article instructions and general publication conditions. It is a pre-submission engineering audit, not editorial approval.

## Automated Checks

| Check | Status |
| --- | --- |
{check_lines}

## Counts

- Abstract words: {abstract_words} (journal limit: 350).
- Keywords: {len(keywords)} ({", ".join(keywords)}).
- Main manuscript words, including references/tables/legends: {_main_word_count(paths["main"])}.
- First-seen figure citations: {figures}.
- First-seen table citations: {tables}.
- First-seen additional-file citations: {supplements}.
- Unresolved author-confirmation/completion placeholders in main manuscript: {placeholders}.

## JTM-Specific Structure Note

The user-requested main heading is **Introduction**. The current Springer Nature JTM Research Article guidance commonly labels the corresponding section **Background**. The structured abstract already uses Background, Methods, Results and Conclusions. Confirm the live Editorial Manager template immediately before submission and rename the main Introduction heading to Background if the submission system requires it.

## Items Requiring Author Action

- Missing manuscript sections: {", ".join(missing_sections) if missing_sections else "none"}.
- Required submission items not yet assembled: {", ".join(missing_required_files) if missing_required_files else "none"}.
- Complete author names, affiliations, corresponding-author details and ORCIDs.
- Obtain or document the institutional determination for the public, de-identified-data ethics statement.
- Confirm consent wording, competing interests, funding and CRediT contributions.
- Verify the three placeholder literature groups in the References section and remove all bracketed author instructions.
- Assemble publication-resolution Figures 1-5 and Additional files 1-7 from the audited source outputs.
- Publish a versioned code release and archived identifier if the authors choose to make the code public at submission.
- Recheck file-format, figure-resolution and any article-processing-charge requirements on the live journal site.

## Official Sources

- Research Article instructions: {JTM_GUIDELINES_URL}
- Conditions of publication: {JTM_CONDITIONS_URL}

## Submission Status

**Not ready for automatic submission.** The scientific draft and formatting framework are generated, but author metadata, declarations, reference verification and final figure/supplement assembly remain required.
"""


RISK_PATTERNS = {
    "causal claim": [
        r"\bcaus(?:e[sd]?|al(?:ly)?)\b",
        r"\bmechanistic proof\b",
    ],
    "clinical readiness": [
        r"\bclinical(?:ly)?[- ]ready\b",
        r"\bready for clinical (?:use|deployment)\b",
        r"\bdemonstrates clinical utility\b",
    ],
    "protein overclaim": [
        r"\bdefinitive proteomic validation\b",
        r"\bconclusive proteomic validation\b",
        r"\ball candidates (?:were )?validated\b",
    ],
    "single-cell overclaim": [
        r"\bsingle[- ]cell validated survival model\b",
        r"\bdirect single[- ]cell survival validation\b",
    ],
    "WSI overclaim": [
        r"\bdigital pathology improved prediction\b",
        r"\bWSI improved prediction\b",
        r"\bpathology improved prediction\b",
    ],
    "excess certainty": [
        r"\bproves?\b",
        r"\bguarantees?\b",
    ],
}


def language_risk_report(root: str | Path = ".") -> str:
    paths = stage9_paths(root)
    scanned = [
        paths["main"],
        paths["title_page"],
        paths["cover_letter"],
        paths["declarations"],
    ]
    hits: list[tuple[str, str, str]] = []
    for path in scanned:
        text = _document_text(path)
        for category, patterns in RISK_PATTERNS.items():
            for pattern in patterns:
                for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                    start = max(0, match.start() - 70)
                    end = min(len(text), match.end() + 70)
                    context = " ".join(text[start:end].split())
                    hits.append((path.name, category, context))
    hit_rows = (
        "\n".join(
            f"| {file_name} | {category} | {context} |"
            for file_name, category, context in hits
        )
        if hits
        else "| None | None | No prohibited positive assertions detected. |"
    )
    return f"""# JTM Language Risk Audit

Generated: {datetime.now().isoformat(timespec="seconds")}

## Result

- Documents scanned: {len(scanned)}.
- High-risk phrase hits: {len(hits)}.
- Automated status: {"REVIEW REQUIRED" if hits else "PASS"}.

| File | Risk category | Context |
| --- | --- | --- |
{hit_rows}

## Required Claim Language

- Use **associated with**, **directionally consistent**, **supports interpretation**, or **exploratory evidence** for biological findings.
- Describe cell states as curated signature scores inferred from bulk RNA-seq.
- Describe HPA as qualitative/IHC-link evidence.
- Describe CPTAC as exploratory quantitative support because only 26 deaths were available.
- State that the RNA increment over clinical-only Cox was modest.
- State that the WSI pilot did not show stable added predictive value.
- Describe the model as a research risk score pending prospective, calibrated evaluation.

## Phrases To Avoid In Positive Assertions

- Causal or mechanistic confirmation.
- A model ready for clinical use or deployment.
- Conclusive protein validation.
- Direct single-cell survival validation.
- A claim that digital pathology improved prediction.

Negated limitations are scientifically appropriate, but this audit deliberately uses conservative pattern matching. Any reported hit requires human review in context.
"""


def evidence_claim_report(root: str | Path = ".") -> str:
    evidence = load_stage8_evidence(root)
    external = ", ".join(
        f"{row['cohort']} {row['c_index']:.3f}" for row in _external_rows(evidence)
    )
    rows = [
        (
            "The frozen score externally ranked survival risk.",
            f"Four GEO C-indices: {external}; continuous-score Cox P<0.01 in each cohort.",
            "Supported",
            "Use risk-ranking language; do not imply calibrated clinical utility.",
        ),
        (
            "RNA added prognostic information beyond clinical variables.",
            f"Nested C-index {evidence.nested_combined_cindex:.3f} versus {evidence.nested_clinical_cindex:.3f}.",
            "Qualified",
            "Call the increment modest; no strong superiority claim.",
        ),
        (
            "A single TCGA cutoff transfers across platforms.",
            "Primary TCGA cutoff failed or collapsed groups in external cohorts.",
            "Not supported",
            "Cohort-median analyses are sensitivity analyses only.",
        ),
        (
            "Higher risk reflects hypoxia, proliferation, EMT-like and CAF programs.",
            "TCGA correlations/group differences and concordant GEO directions.",
            "Supported as association",
            "Bulk signature inference; no causality or direct cell counting.",
        ),
        (
            "Cell-state analysis is direct single-cell outcome validation.",
            "No raw single-cell cohort with linked survival was analysed.",
            "Not supported",
            "Use curated cell-state signatures inferred from bulk RNA.",
        ),
        (
            "HPA validates protein expression quantitatively.",
            f"{evidence.hpa_supported_count}/{evidence.hpa_candidate_count} records/links; no uniform structured IHC endpoint.",
            "Not supported",
            "HPA is qualitative/IHC-link evidence.",
        ),
        (
            "CPTAC provides orthogonal protein support.",
            f"{evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} candidates; LDHA/MKI67/CDK1 stage support; hypoxia module FDR={evidence.cptac_hypoxia_fdr:.4f}.",
            "Exploratory support",
            f"Only {evidence.cptac_deaths} deaths; no broad or conclusive validation claim.",
        ),
        (
            "WSI improved prediction.",
            f"Pilot test C-index clinical {evidence.wsi_clinical_cindex:.3f}, pathology {evidence.wsi_pathology_cindex:.3f}, fusion {evidence.wsi_fusion_cindex:.3f}.",
            "Not supported",
            "Report technical feasibility and neutral predictive result.",
        ),
        (
            "The model can be used in current clinical care.",
            "Retrospective public cohorts; no prospective utility, portable cutoff or decision-impact study.",
            "Not supported",
            "Require prospective assay-harmonized validation.",
        ),
    ]
    table = "\n".join(
        f"| {claim} | {evidence_text} | {status} | {permitted} |"
        for claim, evidence_text, status, permitted in rows
    )
    return f"""# JTM Evidence-Claim Audit

Generated: {datetime.now().isoformat(timespec="seconds")}

## Evidence Hierarchy

1. Primary prediction evidence: leakage-controlled TCGA nested cross-validation and locked GEO validation.
2. Biological interpretation: prespecified bulk RNA cell-state associations with GEO direction checks.
3. Orthogonal support: qualitative HPA links and exploratory CPTAC proteomics.
4. Supplementary neutral evidence: WSI technical feasibility without stable added predictive value.

## Claim Map

| Proposed claim | Direct project evidence | Status | Permitted manuscript wording |
| --- | --- | --- | --- |
{table}

## Integrity Decision

The manuscript may lead with externally reproducible **risk ranking** and a convergent **association** with hypoxia, proliferation, EMT-like and stromal programs. It must not lead with multimodal superiority, clinical deployment, direct single-cell validation, conclusive protein validation or pathology improvement.
"""


def write_compliance_outputs(root: str | Path = ".") -> dict[str, Path]:
    paths = stage9_paths(root)
    paths["reports"].mkdir(parents=True, exist_ok=True)
    build_submission_checklist(root)
    paths["compliance"].write_text(compliance_report(root), encoding="utf-8")
    paths["claims"].write_text(evidence_claim_report(root), encoding="utf-8")
    return {
        "checklist": paths["checklist"],
        "compliance": paths["compliance"],
        "claims": paths["claims"],
    }


def write_language_audit(root: str | Path = ".") -> Path:
    paths = stage9_paths(root)
    paths["reports"].mkdir(parents=True, exist_ok=True)
    paths["language"].write_text(language_risk_report(root), encoding="utf-8")
    return paths["language"]


def append_stage9_audit(root: str | Path, message: str) -> Path:
    path = stage9_paths(root)["audit"]
    with path.open("a", encoding="utf-8") as handle:
        handle.write(message)
    return path
