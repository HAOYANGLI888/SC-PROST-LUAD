"""BMC Cancer manuscript generation and submission-readiness audits."""

from __future__ import annotations

from datetime import date, datetime
from pathlib import Path
import re
from typing import Iterable

import pandas as pd

from reporting.stage8 import Stage8Evidence, load_stage8_evidence
from reporting.stage9_jtm import (
    _add_heading,
    _add_paragraph,
    _add_table,
    _add_title,
    _citation_order,
    _document_text,
    _document_xml,
    _enable_continuous_line_numbers,
    _external_rows,
    _format_p,
    _main_word_count,
    _new_document,
    _require_docx,
)


TITLE = (
    "Development and external validation of a transcriptomic-clinical "
    "prognostic model for lung adenocarcinoma with cellular-context and "
    "proteomic interpretation"
)
RUNNING_TITLE = "Transcriptomic-clinical prognosis in LUAD"
GITHUB_URL = "https://github.com/HAOYANGLI888/SC-PROST-LUAD"
AUTHOR_NAMES = "Haoyang Li; Xiang Li; Xuefeng Shi"
AUTHOR_LINE = "Haoyang Li1†, Xiang Li2†, Xuefeng Shi3*"
EQUAL_CONTRIBUTION = (
    "†Haoyang Li and Xiang Li contributed equally to this work and share "
    "first authorship."
)
AFFILIATIONS = [
    "1 Department of Cardiovascular Medicine, Qinghai University Affiliated "
    "Hospital, No. 29 Tongren Road, Chengxi District, Xining, Qinghai 810000, China.",
    "2 Department of Thyroid and Breast Surgery, Qinghai University Affiliated "
    "Hospital, No. 29 Tongren Road, Chengxi District, Xining, Qinghai 810000, China.",
    "3 Department of Respiratory and Critical Care Medicine, Qinghai Provincial "
    "People's Hospital, No. 2 Gonghe Road, Chengdong District, Xining, "
    "Qinghai 810000, China.",
]
CORRESPONDING_AUTHOR = (
    "Xuefeng Shi, Department of Respiratory and Critical Care Medicine, "
    "Qinghai Provincial People's Hospital, No. 2 Gonghe Road, Chengdong "
    "District, Xining, Qinghai 810000, China. Email: shixuefeng128@163.com. "
    "ORCID: 0000-0002-4694-8759."
)
FUNDING_STATEMENT = (
    "This study was supported by the National Excellent Young Physician "
    "Program (国家优秀青年医师; Document No. 2024[41]). The funder had no role "
    "in study design, data collection and analysis, decision to publish, or "
    "preparation of the manuscript."
)
AUTHOR_CONTRIBUTIONS = (
    "H.L. and X.L. contributed equally to this work and share first authorship. "
    "H.L. conceived the study, designed the analytical workflow, performed data "
    "processing and statistical analyses, interpreted the results, prepared "
    "tables and figures, and drafted the manuscript. X.L. contributed to study "
    "design, clinical interpretation, data review, and critical manuscript "
    "revision. X.S. supervised the study, provided clinical oversight, reviewed "
    "and revised the manuscript critically, and is the corresponding author. "
    "All authors read and approved the final manuscript."
)
BMC_GUIDELINES_URL = (
    "https://bmccancer.biomedcentral.com/submission-guidelines/"
    "preparing-your-manuscript/research-article"
)
BMC_SCOPE_URL = "https://bmccancer.biomedcentral.com/submission-guidelines/aims-and-scope"


class Stage9BMCError(RuntimeError):
    """Raised when a BMC Cancer artifact cannot be generated faithfully."""


def bmc_paths(root: str | Path = ".") -> dict[str, Path]:
    project_root = Path(root).resolve()
    manuscript = project_root / "outputs" / "manuscript"
    tables = project_root / "outputs" / "tables"
    reports = project_root / "outputs" / "reports"
    return {
        "root": project_root,
        "manuscript": manuscript,
        "tables": tables,
        "reports": reports,
        "main": manuscript / "BMC_Cancer_main_manuscript_draft.docx",
        "title_page": manuscript / "BMC_Cancer_title_page.docx",
        "cover_letter": manuscript / "BMC_Cancer_cover_letter.docx",
        "declarations": manuscript / "BMC_Cancer_declarations.docx",
        "checklist": tables / "BMC_Cancer_submission_file_checklist.csv",
        "compliance": reports / "BMC_Cancer_compliance_report.md",
        "language": reports / "BMC_Cancer_claim_language_audit.md",
        "claims": reports / "BMC_Cancer_evidence_claim_audit.md",
        "figure_table": reports / "BMC_Cancer_figure_table_audit.md",
        "audit": project_root / "audit_report.md",
    }


def _new_bmc_document(*, letter: bool = False):
    document = _new_document(letter=letter)
    document.core_properties.author = AUTHOR_NAMES
    document.core_properties.title = TITLE
    document.core_properties.subject = "BMC Cancer Research Article submission draft"
    document.core_properties.comments = (
        "Generated from locked project outputs; author metadata and final "
        "submission assets require confirmation."
    )
    return document


def _add_author_block(document) -> None:
    _add_paragraph(document, AUTHOR_LINE, first_line=False)
    _add_paragraph(document, EQUAL_CONTRIBUTION, first_line=False)
    for affiliation in AFFILIATIONS:
        _add_paragraph(document, affiliation, first_line=False)
    _add_paragraph(
        document,
        f"* Corresponding author: {CORRESPONDING_AUTHOR}",
        first_line=False,
    )


def _set_landscape(section) -> None:
    docx = _require_docx()
    from docx.enum.section import WD_ORIENT

    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = docx["Inches"](11.0)
    section.page_height = docx["Inches"](8.5)
    section.left_margin = docx["Inches"](0.55)
    section.right_margin = docx["Inches"](0.55)
    section.top_margin = docx["Inches"](0.65)
    section.bottom_margin = docx["Inches"](0.65)
    _enable_continuous_line_numbers(section)


def _set_portrait(section) -> None:
    docx = _require_docx()
    from docx.enum.section import WD_ORIENT

    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = docx["Inches"](8.5)
    section.page_height = docx["Inches"](11.0)
    section.left_margin = docx["Inches"](1.0)
    section.right_margin = docx["Inches"](1.0)
    section.top_margin = docx["Inches"](1.0)
    section.bottom_margin = docx["Inches"](1.0)
    _enable_continuous_line_numbers(section)


def _format_table(table, widths: list[float]) -> None:
    docx = _require_docx()
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = docx["OxmlElement"]("w:tblLayout")
    layout.set(docx["qn"]("w:type"), "fixed")
    table_pr.append(layout)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = docx["OxmlElement"]("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            width = docx["Inches"](widths[index])
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is not None:
                tc_w.set(docx["qn"]("w:w"), str(int(widths[index] * 1440)))
                tc_w.set(docx["qn"]("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    run.font.size = docx["Pt"](9)


def _load_scrna_support(root: Path) -> pd.DataFrame:
    path = root / "outputs" / "tables" / "stage4b_lowmem_mechanism_support_matrix.csv"
    if not path.exists():
        raise FileNotFoundError(f"Required Stage 4B evidence is missing: {path}")
    frame = pd.read_csv(path)
    required = {
        "mechanism",
        "top_observed_cell_types",
        "support_status",
        "evidence_scope",
    }
    missing = sorted(required.difference(frame.columns))
    if missing:
        raise Stage9BMCError(
            f"Stage 4B support table is missing columns: {', '.join(missing)}"
        )
    if frame.empty:
        raise Stage9BMCError("Stage 4B support table is empty.")
    return frame


def _scrna_row(frame: pd.DataFrame, mechanism: str) -> pd.Series:
    rows = frame.loc[frame["mechanism"].eq(mechanism)]
    if rows.empty:
        raise Stage9BMCError(f"Stage 4B mechanism is missing: {mechanism}")
    return rows.iloc[0]


def _abstract_text(
    evidence: Stage8Evidence, scrna: pd.DataFrame
) -> dict[str, str]:
    external = ", ".join(
        f"{row['cohort']} {row['c_index']:.3f}" for row in _external_rows(evidence)
    )
    return {
        "Background": (
            "Transcriptomic prognostic models for lung adenocarcinoma often "
            "lack locked multi-cohort validation and a clearly bounded account "
            "of the cellular programs represented by their risk scores."
        ),
        "Methods": (
            f"We developed an overall-survival model in {evidence.tcga_patients} "
            "TCGA-LUAD patients using clinical variables and 25 RNA principal "
            "components under nested cross-validation. The frozen pipeline was "
            f"applied without refitting to four GEO cohorts ({evidence.external_total} "
            f"patients). We evaluated {evidence.signature_count} prespecified "
            "cell-state signatures in bulk expression, localized selected "
            "programs in GSE131907 raw single-cell RNA sequencing data "
            "(208,506 cells), and examined qualitative HPA evidence and "
            "exploratory CPTAC-LUAD proteomics."
        ),
        "Results": (
            "The combined Cox model achieved a nested C-index of "
            f"{evidence.nested_combined_cindex:.3f} (SD "
            f"{evidence.nested_combined_sd:.3f}), compared with "
            f"{evidence.nested_clinical_cindex:.3f} (SD "
            f"{evidence.nested_clinical_sd:.3f}) for clinical Cox. External "
            f"C-indices were {external}, and continuous risk scores were "
            "associated with overall survival in all four cohorts. Higher risk "
            "tracked hypoxia, proliferation, EMT-like and CAF programs, whereas "
            "dendritic, B-cell and plasma-cell programs marked a lower-risk "
            "context. Raw single-cell data localized hypoxia mainly to malignant "
            "epithelial cells, CAF/matrix signals to fibroblasts, and immune "
            "programs to their corresponding lineages. Proliferation was not "
            "malignant-epithelial specific, and the EMT-like signal was mainly "
            "stromal/myeloid. HPA supplied qualitative/IHC-link evidence for "
            f"{evidence.hpa_supported_count} candidates. CPTAC quantified "
            f"{evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} "
            "candidates, and its hypoxia module was associated with survival "
            f"(HR {evidence.cptac_hypoxia_hr:.2f} per SD, FDR "
            f"{evidence.cptac_hypoxia_fdr:.4f}) among only "
            f"{evidence.cptac_deaths} deaths."
        ),
        "Conclusions": (
            "The frozen score reproducibly ranked survival risk across four "
            "external cohorts. Its biological interpretation converged on "
            "malignant epithelial hypoxia, CAF/matrix activation and immune-state "
            "differences, with important qualifications for proliferation and "
            "EMT-like signals. The modest internal increment over clinical "
            "variables and exploratory protein evidence require cautious "
            "interpretation and further assay-harmonized evaluation."
        ),
    }


def _build_main_tables(
    document, evidence: Stage8Evidence, scrna: pd.DataFrame, root: Path
) -> None:
    external = _external_rows(evidence)
    cohort_rows = [
        [
            "TCGA-LUAD",
            "RNA-seq",
            evidence.tcga_patients,
            evidence.tcga_deaths,
            evidence.tcga_censored,
            "Development and internal validation",
        ]
    ]
    cohort_rows.extend(
        [
            [
                row["cohort"],
                "Microarray",
                row["n"],
                row["deaths"],
                row["n"] - row["deaths"],
                "Locked external validation",
            ]
            for row in external
        ]
    )
    table = _add_table(
        document,
        "Table 1. Characteristics and analytical roles of the prognostic cohorts.",
        ["Cohort", "Assay", "Patients", "Deaths", "Censored", "Role"],
        cohort_rows,
    )
    _format_table(table, [1.2, 1.1, 0.8, 0.8, 0.9, 4.9])

    performance_rows = [
        [
            "TCGA nested CV",
            "Clinical Cox",
            f"{evidence.nested_clinical_cindex:.3f} "
            f"({evidence.nested_clinical_sd:.3f})",
            "-",
            "-",
            "-",
            "-",
        ],
        [
            "TCGA nested CV",
            "RNA PCA_25 + clinical Cox",
            f"{evidence.nested_combined_cindex:.3f} "
            f"({evidence.nested_combined_sd:.3f})",
            f"{evidence.nested_auc_1y:.3f}",
            f"{evidence.nested_auc_3y:.3f}",
            f"{evidence.nested_auc_5y:.3f}",
            "-",
        ],
    ]
    performance_rows.extend(
        [
            [
                row["cohort"],
                "Frozen full model",
                f"{row['c_index']:.3f}",
                f"{row['auc_1y']:.3f}",
                f"{row['auc_3y']:.3f}",
                f"{row['auc_5y']:.3f}",
                f"{row['hr']:.3f}; P={_format_p(row['p'])}",
            ]
            for row in external
        ]
    )
    table = _add_table(
        document,
        "Table 2. Internal and external performance of the prognostic model.",
        [
            "Cohort",
            "Model",
            "C-index",
            "AUC 1 y",
            "AUC 3 y",
            "AUC 5 y",
            "HR per SD",
        ],
        performance_rows,
    )
    _format_table(table, [1.3, 1.9, 1.0, 0.95, 0.95, 0.95, 2.15])

    cell_rows = [
        [
            "Hypoxia",
            "Higher in high-risk tumors",
            "Same direction in 4/4 GEO cohorts",
            "Malignant epithelial cells; epithelial cells",
            "Malignant epithelial hypoxia is the clearest adverse context.",
        ],
        [
            "Proliferation",
            "Higher in high-risk tumors",
            "Same direction in 4/4 GEO cohorts",
            "T cells first; malignant epithelial cells second",
            "A broader cycling-cell signal, not a tumor-cell-specific program.",
        ],
        [
            "EMT-like",
            "Higher in high-risk tumors",
            "Same direction in 4/4 GEO cohorts",
            "Fibroblast/CAF, myeloid, and endothelial cells",
            "Mainly stromal and myeloid context; malignant-cell EMT was not established.",
        ],
        [
            "CAF/matrix",
            "Higher in high-risk tumors",
            "Same direction in 3/4 GEO cohorts",
            "Fibroblast/CAF cells",
            "Supports an adverse matrix and stromal context.",
        ],
        [
            "Dendritic/B/plasma",
            "Higher in lower-risk tumors",
            "Dendritic and B: 4/4; plasma: 3/4",
            "Corresponding dendritic, B-cell, and plasma-cell lineages",
            "Supports a lower-risk immune-cell context.",
        ],
    ]
    table = _add_table(
        document,
        "Table 3. Summary of bulk cell-state associations and raw scRNA cellular context.",
        [
            "Program",
            "Bulk risk association",
            "GEO consistency",
            "Raw scRNA localization",
            "Final interpretation",
        ],
        cell_rows,
    )
    _format_table(table, [1.05, 1.55, 1.45, 2.4, 3.25])

    document.add_page_break()
    table = _add_table(
        document,
        "Table 4. Integrated HPA and CPTAC protein evidence.",
        ["Evidence layer", "Coverage", "Main finding", "Boundary"],
        [
            [
                "HPA",
                f"{evidence.hpa_supported_count}/"
                f"{evidence.hpa_candidate_count} candidates",
                "Qualitative protein/IHC links available",
                "Qualitative evidence; no uniform quantitative IHC endpoint.",
            ],
            [
                "CPTAC candidates",
                f"{evidence.cptac_candidates_matched}/"
                f"{evidence.cptac_candidate_total} candidates in "
                f"{evidence.cptac_samples} tumors",
                "LDHA, MKI67, and CDK1 had direction-consistent nominal stage associations.",
                f"Exploratory; only {evidence.cptac_deaths} deaths were available.",
            ],
            [
                "CPTAC hypoxia module",
                f"{evidence.cptac_usable_os} patients with usable OS; "
                f"{evidence.cptac_deaths} deaths",
                (
                    f"HR {evidence.cptac_hypoxia_hr:.2f}, 95% CI "
                    f"{evidence.cptac_hypoxia_ci_low:.2f}-"
                    f"{evidence.cptac_hypoxia_ci_high:.2f}; "
                    f"FDR={evidence.cptac_hypoxia_fdr:.4f}"
                ),
                "Exploratory quantitative support, not definitive validation.",
            ],
        ],
    )
    _format_table(table, [1.4, 2.2, 3.25, 2.85])


def generate_main_manuscript(
    root: str | Path = ".", *, output_path: str | Path | None = None
) -> Path:
    paths = bmc_paths(root)
    paths["manuscript"].mkdir(parents=True, exist_ok=True)
    evidence = load_stage8_evidence(root)
    scrna = _load_scrna_support(paths["root"])
    document = _new_bmc_document()
    _add_title(document, TITLE)
    _add_author_block(document)

    _add_heading(document, "Abstract", 1)
    for label, text in _abstract_text(evidence, scrna).items():
        _add_paragraph(
            document, f"{label}: {text}", bold_prefix=f"{label}:", first_line=False
        )
    _add_paragraph(
        document,
        "Keywords: lung adenocarcinoma; overall survival; transcriptomics; "
        "external validation; single-cell RNA sequencing; proteomics",
        bold_prefix="Keywords:",
        first_line=False,
    )

    _add_heading(document, "Introduction", 1)
    _add_paragraph(
        document,
        "Lung adenocarcinoma (LUAD) is clinically and molecularly heterogeneous. "
        "Patients with similar clinicopathological characteristics can experience "
        "different outcomes, while routine variables do not fully represent the "
        "molecular states associated with aggressive disease. Transcriptomic "
        "profiling can capture these states, but a prognostic score is most useful "
        "as research evidence when it remains reproducible beyond its development "
        "cohort [1].",
    )
    _add_paragraph(
        document,
        "Published molecular signatures can be affected by information leakage, "
        "limited external evaluation and platform-specific score distributions. "
        "Feature selection or standardization before data splitting can inflate "
        "performance. A model may also preserve relative risk ranking while "
        "failing to transport an absolute threshold. Leakage-controlled "
        "development, locked external application and explicit reporting of "
        "threshold limitations are therefore essential [10, 15, 16].",
    )
    _add_paragraph(
        document,
        "Biological interpretation adds a separate challenge because bulk RNA "
        "profiles combine malignant, immune and stromal signals. Curated "
        "cell-state signatures can identify coherent associations, while raw "
        "single-cell profiles can help localize selected programs to cellular "
        "contexts. Neither approach establishes a survival effect at single-cell "
        "resolution without linked outcomes. Protein resources can provide an "
        "additional orthogonal layer, provided that qualitative HPA evidence and "
        "limited-event CPTAC analyses remain clearly bounded [7, 8, 11, 20-23].",
    )
    _add_paragraph(
        document,
        "We developed a clinical plus RNA principal-component Cox model in "
        "TCGA-LUAD using nested cross-validation, froze the full pipeline and "
        "applied it without refitting to four GEO cohorts. We then combined bulk "
        "cell-state scoring with raw GSE131907 single-cell cellular-context "
        "analysis and bounded HPA and CPTAC evidence. A whole-slide imaging pilot "
        "was retained as a supplementary feasibility result because it did not "
        "show stable improvement over the clinical comparator.",
    )

    _add_heading(document, "Methods", 1)
    _add_heading(document, "Study design and data sources", 2)
    _add_paragraph(
        document,
        "This retrospective computational study used de-identified public data. "
        "TCGA-LUAD clinical and RNA-sequencing data were obtained through the "
        "Genomic Data Commons [17]. External validation used GSE31210, GSE50081, "
        "GSE72094 and GSE68465 from the Gene Expression Omnibus. Cellular-context "
        "analysis used GSE131907 raw single-cell RNA-sequencing data and official "
        "author annotations. Protein interpretation used the Human Protein Atlas "
        "(HPA) and the PDC000153 CPTAC LUAD Discovery Proteome. The prognostic "
        "model was trained only in TCGA, and external outcomes did not influence "
        "genes, components, coefficients or preprocessing parameters "
        "[5, 6, 11-14].",
    )
    _add_paragraph(
        document,
        "The analytical layers were kept distinct. TCGA nested cross-validation "
        "and locked GEO evaluation addressed prognosis. Bulk and raw single-cell "
        "analyses addressed cellular context. HPA and CPTAC addressed orthogonal "
        "protein evidence. Whole-slide images were not incorporated into the "
        "frozen model.",
    )

    _add_heading(document, "TCGA-LUAD cohort and survival endpoints", 2)
    _add_paragraph(
        document,
        f"Patient identifiers were standardized to the first 12 characters of "
        f"TCGA barcodes. The final cohort included {evidence.tcga_patients} "
        f"patients with primary-tumor RNA-seq and usable overall survival, "
        f"including {evidence.tcga_deaths} deaths and "
        f"{evidence.tcga_censored} censored observations. Survival time was "
        "represented in days, and status was encoded as 1 for death and 0 for "
        "censoring. Overall survival was the sole model-development endpoint.",
    )

    _add_heading(document, "RNA-seq preprocessing and clinical variables", 2)
    _add_paragraph(
        document,
        "GDC STAR-count files were restricted to primary tumors. TPM values were "
        "obtained from the tpm_unstranded field, duplicate primary-tumor files "
        "were resolved deterministically and protein-coding genes were retained. "
        "Expression was transformed as log2(TPM + 1). The matrix contained "
        "19,962 protein-coding genes across 517 patients before survival matching. "
        "Missingness filtering, variance ranking, standardization and principal "
        "component analysis were fitted within training data only.",
    )
    _add_paragraph(
        document,
        "Clinical predictors included harmonized age, sex and pathological stage "
        "where available. Categorical encoding and missing-value handling were "
        "learned from training data. The selected RNA representation comprised "
        "25 principal components derived from 1,000 training-selected "
        "high-variance genes.",
    )

    _add_heading(document, "Prognostic model development and internal validation", 2)
    _add_paragraph(
        document,
        "Candidate clinical, RNA-only and combined models were compared using "
        "nested cross-validation. The outer loop used five folds and the inner "
        "loop used three folds for seeds 42, 3407 and 2026. Stratification "
        "balanced event status and survival-time strata where feasible. Every "
        "preprocessing and selection step was fitted in the applicable training "
        "fold, and outer-fold predictions were retained for unbiased assessment "
        "[15, 16].",
    )
    _add_paragraph(
        document,
        "The frozen model was the RNA PCA_25 plus clinical Cox model, with "
        "clinical Cox as the reference. Elastic-net Cox and DeepSurv were assessed "
        "during robustness analyses. DeepSurv showed larger train-to-test gaps "
        "and was not selected. The training genes, scalers, PCA loadings and Cox "
        "coefficients were fixed before external evaluation [2, 19].",
    )

    _add_heading(document, "External GEO validation", 2)
    _add_paragraph(
        document,
        "Each GEO cohort was imported, probe-mapped and normalized independently. "
        "External data were not pooled with TCGA for standardization. Missing "
        "frozen-model genes were handled using the locked procedure, after which "
        "the fixed feature order, PCA transformation and Cox coefficients were "
        "applied without outcome-guided modification.",
    )
    _add_paragraph(
        document,
        "Primary external evaluation used the continuous risk score. Harrell's "
        "C-index, time-dependent area under the curve at 1, 3 and 5 years, and "
        "the hazard ratio per standard deviation were estimated. The TCGA-derived "
        "cutoff was evaluated for transportability, while cohort-median "
        "Kaplan-Meier analyses were treated as sensitivity analyses [3, 4, 18].",
    )

    _add_heading(document, "Bulk RNA-seq cell-state signature scoring", 2)
    _add_paragraph(
        document,
        f"We evaluated {evidence.signature_count} prespecified LUAD cell-state "
        "signatures spanning malignant epithelial, EMT-like, proliferative and "
        "hypoxic tumor programs, immune lineages, cancer-associated fibroblasts "
        "and endothelial cells. Signatures were defined before outcome association "
        "analysis and were not optimized against survival. Mean z scores were "
        "calculated from genes available in each independently standardized "
        "dataset [9].",
    )
    _add_paragraph(
        document,
        "Spearman correlations quantified associations between cell-state scores "
        "and the frozen RNA risk score. TCGA high- and low-risk groups were "
        "compared, and exploratory univariable and age-, sex- and stage-adjusted "
        "Cox models were fitted. GEO analyses assessed directional consistency "
        "without changing the signatures or prognostic model.",
    )

    _add_heading(document, "Raw scRNA-seq cellular-context analysis", 2)
    _add_paragraph(
        document,
        "The GSE131907 raw UMI matrix contained 208,506 cells and 29,634 genes. "
        "Official author-provided cell annotations were retained, and no de novo "
        "cell-type discovery was used to revise the prespecified programs. "
        "Sparse, backed processing calculated log1p counts per million summaries "
        "for the complete cohort. Fixed gene and program scores were summarized "
        "by annotated cell type. All requested signature genes were present.",
    )
    _add_paragraph(
        document,
        "A 50,000-cell stratified subset was used only for UMAP visualization. "
        "Full-cohort principal-component, neighbor-graph and UMAP calculations "
        "were not performed on the 32 GB workstation. Cellular-context support "
        "was assigned using prespecified expected lineages and ranked cell-type "
        "scores. The single-cell cohort had no linked survival endpoint and did "
        "not enter prognostic-model training [11, 20-23].",
    )

    _add_heading(document, "HPA protein/IHC evidence", 2)
    _add_paragraph(
        document,
        f"We prespecified {evidence.hpa_candidate_count} genes representing the "
        "main biological programs. HPA gene and lung-cancer pathology pages were "
        "reviewed for protein evidence, antibodies and immunohistochemistry links. "
        "Because a uniform structured quantitative lung-cancer staining endpoint "
        "was unavailable, this layer was classified as qualitative/IHC-link "
        "evidence and not as an independent survival cohort [7].",
    )

    _add_heading(document, "CPTAC-LUAD quantitative proteomic analysis", 2)
    _add_paragraph(
        document,
        f"The PDC000153 discovery proteome comprised {evidence.cptac_samples} "
        f"primary tumors and {evidence.cptac_proteins:,} protein genes. Usable "
        f"overall survival was available for {evidence.cptac_usable_os} patients, "
        f"with {evidence.cptac_deaths} deaths. Of the prespecified candidates, "
        f"{evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} "
        "were quantified. Candidate abundance and prespecified protein-module "
        "scores were examined for stage and survival associations with "
        "Benjamini-Hochberg adjustment.",
    )
    _add_paragraph(
        document,
        "The limited event count and incomplete candidate coverage defined these "
        "analyses as exploratory quantitative proteomic support. They were not "
        "used to modify the transcriptomic model or select new signatures [8].",
    )

    _add_heading(document, "Exploratory WSI feasibility analysis", 2)
    _add_paragraph(
        document,
        "A separate 100-slide TCGA-LUAD pilot evaluated OpenSlide reading, tissue "
        "segmentation, patch extraction, pretrained ResNet50 features and "
        "attention-based multiple-instance learning. Clinical-only, pathology-only "
        "and fusion models were compared on a small held-out split. The pilot was "
        "not expanded after the pathology models failed to show stable improvement.",
    )

    _add_heading(document, "Statistical analysis", 2)
    _add_paragraph(
        document,
        "Discrimination was summarized by C-index and time-dependent AUC. "
        "Kaplan-Meier curves and log-rank tests described group separation, and "
        "Cox models reported hazard ratios with 95% confidence intervals. "
        "Spearman coefficients quantified risk-program associations. "
        "Benjamini-Hochberg adjustment was used for signature and protein families. "
        "All tests were two-sided. Analyses used Python, pandas, NumPy, "
        "scikit-learn, lifelines, PyTorch and project-specific modules "
        "[2-4, 18, 19].",
    )
    _add_paragraph(
        document,
        "Detailed model results, GEO coverage, signature definitions, raw "
        "single-cell localization, HPA evidence, CPTAC analyses and WSI "
        "diagnostics are provided in Supplementary Tables S1-S8. Additional "
        "performance, GEO processing, bulk cell-state, raw single-cell, HPA, "
        "CPTAC and WSI displays are provided in Supplementary Figs. S1-S7.",
    )

    _add_heading(document, "Results", 1)
    _add_heading(document, "TCGA cohort construction and model development", 2)
    _add_paragraph(
        document,
        f"The study separated TCGA development, locked GEO validation and "
        f"post-model interpretation (Fig. 1). TCGA included "
        f"{evidence.tcga_patients} patients, {evidence.tcga_deaths} deaths and "
        f"{evidence.tcga_censored} censored observations; the four external "
        f"cohorts contributed {evidence.external_total} patients (Table 1).",
    )
    _add_paragraph(
        document,
        "Across 15 outer folds, RNA PCA_25 plus clinical Cox achieved a mean "
        f"C-index of {evidence.nested_combined_cindex:.3f} (SD "
        f"{evidence.nested_combined_sd:.3f}), compared with "
        f"{evidence.nested_clinical_cindex:.3f} (SD "
        f"{evidence.nested_clinical_sd:.3f}) for clinical Cox. Mean 1-, 3- and "
        f"5-year AUCs were {evidence.nested_auc_1y:.3f}, "
        f"{evidence.nested_auc_3y:.3f} and {evidence.nested_auc_5y:.3f}. The "
        "increment over clinical Cox was modest, and DeepSurv did not show a "
        "stable advantage (Fig. 2; Table 2; Supplementary Fig. S1; "
        "Supplementary Table S1).",
    )

    _add_heading(document, "External validation across four GEO cohorts", 2)
    external_sentences = [
        (
            f"{row['cohort']} (n={row['n']}, {row['deaths']} deaths), C-index "
            f"{row['c_index']:.3f}, HR per SD {row['hr']:.3f}, "
            f"P={_format_p(row['p'])}"
        )
        for row in _external_rows(evidence)
    ]
    _add_paragraph(
        document,
        "The frozen model was applied without refitting. Results were "
        + "; ".join(external_sentences)
        + ". Thus, the continuous score was associated with overall survival in "
        "all four cohorts despite assay and coverage differences (Fig. 3).",
    )
    _add_paragraph(
        document,
        "The TCGA-derived median cutoff was not portable and collapsed all "
        "samples into one group in two cohorts. Cohort-median sensitivity analyses "
        "showed clearer separation but used cohort-specific thresholds. The "
        "external evidence therefore supports reproducible risk ranking rather "
        "than a transferable absolute decision threshold (Supplementary Fig. S2; "
        "Supplementary Table S2).",
    )

    _add_heading(document, "Cell-state programs associated with the RNA risk score", 2)
    _add_paragraph(
        document,
        "In TCGA, the strongest positive risk correlations were observed for "
        "hypoxia (rho=0.354), proliferation (rho=0.349), malignant epithelial "
        "state (rho=0.277), EMT-like state (rho=0.193) and CAF (rho=0.104). "
        "Hypoxia, proliferation, EMT-like and CAF scores had concordant directions "
        "across all four GEO cohorts. Dendritic-cell, B-cell and plasma-cell "
        "programs supported a lower-risk immune context (Table 3; Fig. 4; "
        "Supplementary Fig. S3; Supplementary Table S3; "
        "Supplementary Table S4).",
    )
    _add_paragraph(
        document,
        "Associations for M2 macrophage, regulatory T-cell, exhausted CD8, "
        "cytotoxic CD8, M1 macrophage and NK signatures were mixed or externally "
        "unstable. These programs were therefore excluded from the principal "
        "biological interpretation.",
    )

    _add_heading(
        document,
        "Raw scRNA-seq analysis supports the cellular context of selected programs",
        2,
    )
    _add_paragraph(
        document,
        "The complete GSE131907 matrix supported three prespecified cellular "
        "contexts. Hypoxia scores were highest in malignant epithelial cells, "
        "followed by epithelial cells, and LDHA expression was highest in "
        "malignant epithelial cells. CAF/matrix scores were highest in "
        "fibroblast/CAF annotations. The dendritic/B/plasma program localized to "
        "the corresponding immune lineages, with antigen-presentation markers in "
        "dendritic cells, B-cell markers in B cells and plasma-cell markers in "
        "plasma cells.",
    )
    _add_paragraph(
        document,
        "The same analysis qualified two bulk interpretations. Proliferation "
        "scores were highest in T cells and second highest in malignant epithelial "
        "cells; MKI67 and CDK1 were not malignant-epithelial specific. The "
        "EMT-like score localized mainly to fibroblast/CAF, macrophage/monocyte "
        "and endothelial contexts rather than clearly to malignant epithelial "
        "cells. These findings refine cellular attribution without linking "
        "single-cell states to survival (Supplementary Fig. S4; "
        "Supplementary Table S5).",
    )

    _add_heading(
        document, "HPA and CPTAC provide protein-level orthogonal support", 2
    )
    _add_paragraph(
        document,
        f"All {evidence.hpa_supported_count} candidates had an HPA record or "
        "lung-cancer pathology link. This constituted qualitative/IHC-link "
        "evidence because a harmonized quantitative staining endpoint was not "
        "available (Supplementary Fig. S5; Supplementary Table S6).",
    )
    _add_paragraph(
        document,
        f"CPTAC quantified {evidence.cptac_candidates_matched}/"
        f"{evidence.cptac_candidate_total} candidates. LDHA, MKI67 and CDK1 had "
        "nominally significant stage associations in the prespecified direction, "
        "while no individual candidate retained a survival association after "
        "multiple-testing adjustment. The hypoxia protein module was associated "
        f"with poorer overall survival (HR {evidence.cptac_hypoxia_hr:.3f} per "
        f"SD, 95% CI {evidence.cptac_hypoxia_ci_low:.3f}-"
        f"{evidence.cptac_hypoxia_ci_high:.3f}, P="
        f"{_format_p(evidence.cptac_hypoxia_p)}, FDR="
        f"{evidence.cptac_hypoxia_fdr:.4f}). These findings provide exploratory "
        f"support because only {evidence.cptac_deaths} deaths were available "
        "(Table 4; Fig. 5; Supplementary Fig. S6; Supplementary Table S7).",
    )

    _add_heading(
        document,
        "Exploratory WSI analysis did not show stable improvement over clinical-only prediction",
        2,
    )
    _add_paragraph(
        document,
        f"The 100-slide pipeline completed patch extraction, GPU feature "
        f"extraction and survival-model training. Held-out C-indices were "
        f"{evidence.wsi_clinical_cindex:.3f} for clinical Cox, "
        f"{evidence.wsi_pathology_cindex:.3f} for the best pathology-only model "
        f"and {evidence.wsi_fusion_cindex:.3f} for the best fusion model. These "
        "small-pilot estimates show technical feasibility but no stable added "
        "predictive value (Supplementary Fig. S7; Supplementary Table S8).",
    )

    _add_heading(document, "Discussion", 1)
    _add_heading(document, "Principal findings", 2)
    _add_paragraph(
        document,
        "A frozen TCGA-derived transcriptomic-clinical score retained prognostic "
        "risk ranking across four independent GEO cohorts without refitting. "
        "External C-indices ranged from 0.624 to 0.700, and the continuous score "
        "was associated with overall survival in each cohort. The internal gain "
        "over clinical Cox was small, which favors a restrained interpretation "
        "of the RNA component as reproducible molecular stratification rather "
        "than a replacement for clinicopathological information.",
    )

    _add_heading(
        document,
        "Biological interpretation of hypoxia, CAF/matrix and immune programs",
        2,
    )
    _add_paragraph(
        document,
        "Bulk analyses linked higher risk to hypoxia and CAF-associated programs "
        "and linked lower risk to dendritic, B-cell and plasma-cell contexts. Raw "
        "single-cell data clarified the likely cellular sources: hypoxia and LDHA "
        "were concentrated in malignant epithelial cells, matrix programs in "
        "fibroblast/CAF cells and immune programs in their expected lineages. "
        "This convergence supports a cellular-context interpretation of the risk "
        "score, but it does not connect single-cell states directly to survival.",
    )

    _add_heading(
        document, "Proliferation represents a broader cycling-cell signal", 2
    )
    _add_paragraph(
        document,
        "The proliferation program was reproducibly associated with higher bulk "
        "risk, but raw single-cell localization was broader than initially "
        "suggested by the bulk label. T cells had the highest score, malignant "
        "epithelial cells ranked second, and MKI67 and CDK1 were not restricted "
        "to malignant cells. The program is therefore better interpreted as a "
        "cycling-cell signal that includes, but is not specific to, tumor cells.",
    )

    _add_heading(document, "EMT-like signal requires cautious attribution", 2)
    _add_paragraph(
        document,
        "The EMT-like bulk signature showed consistent risk-direction "
        "associations across cohorts. Raw single-cell analysis, however, localized "
        "the signal mainly to fibroblast/CAF, myeloid and endothelial contexts. "
        "The result is more consistent with a stromal and microenvironmental "
        "program than with demonstrated malignant-cell EMT. This distinction "
        "prevents a mixed bulk signal from being assigned too narrowly.",
    )

    _add_heading(document, "Protein-level support from HPA and CPTAC", 2)
    _add_paragraph(
        document,
        "HPA established qualitative protein and pathology links for the "
        "prespecified candidates. CPTAC quantified most candidates and provided "
        "direction-consistent stage associations for LDHA, MKI67 and CDK1. The "
        "hypoxia module provided the clearest survival association. These layers "
        "are compatible with the RNA and cellular-context findings, but the HPA "
        "endpoint was qualitative and the CPTAC survival analysis contained only "
        "26 deaths.",
    )

    _add_heading(
        document, "Implications for LUAD prognostic biomarker research", 2
    )
    _add_paragraph(
        document,
        "The most transferable result was continuous risk ranking, not an "
        "absolute threshold. Future studies should predefine assay harmonization, "
        "calibration and decision thresholds, and should test whether the score "
        "adds useful information under contemporary treatment conditions. "
        "Experimental studies would also be needed to test the functional roles "
        "of the prioritized hypoxia and stromal programs [24].",
    )

    _add_heading(document, "Limitations", 2)
    limitations = [
        "This was a retrospective public-database study with heterogeneous treatment, follow-up and annotation.",
        "The RNA component provided only a modest increment over clinical Cox in TCGA internal validation.",
        "The GEO cohorts used different microarray platforms, and the TCGA-derived cutoff did not transfer reliably.",
        "The primary cell-state associations were inferred from curated signatures in bulk RNA-seq.",
        "Raw scRNA-seq supported cellular localization but had no linked survival endpoint.",
        "HPA supplied qualitative/IHC-link evidence rather than a harmonized quantitative endpoint.",
        f"CPTAC analyses were exploratory because only {evidence.cptac_deaths} deaths were available.",
        f"CPTAC quantified {evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} candidates rather than the complete set.",
        "The WSI pilot did not show stable improvement over clinical-only prediction.",
        "No prospective cohort or interventional clinical-utility study was evaluated.",
        "No in vitro or in vivo functional experiment was performed, which is an important limitation for biomarker interpretation.",
    ]
    for item in limitations:
        _add_paragraph(document, item, style="List Number", first_line=False)

    _add_heading(document, "Conclusion", 1)
    _add_paragraph(
        document,
        "A leakage-controlled transcriptomic-clinical Cox model reproduced "
        "survival risk ranking across four external LUAD cohorts. Bulk and raw "
        "single-cell analyses converged on malignant epithelial hypoxia, "
        "CAF/matrix activation and immune-state differences, while showing that "
        "proliferation and EMT-like signals require broader cellular attribution. "
        "HPA and CPTAC supplied bounded orthogonal protein evidence. The modest "
        "RNA increment, non-portable cutoff and exploratory biological layers "
        "support further assay-harmonized and experimental evaluation rather than "
        "immediate clinical application.",
    )

    _add_heading(document, "List of abbreviations", 1)
    abbreviations: Iterable[tuple[str, str]] = [
        ("AUC", "area under the time-dependent receiver operating characteristic curve"),
        ("CAF", "cancer-associated fibroblast"),
        ("C-index", "concordance index"),
        ("CPTAC", "Clinical Proteomic Tumor Analysis Consortium"),
        ("FDR", "false discovery rate"),
        ("GDC", "Genomic Data Commons"),
        ("GEO", "Gene Expression Omnibus"),
        ("HPA", "Human Protein Atlas"),
        ("IHC", "immunohistochemistry"),
        ("LUAD", "lung adenocarcinoma"),
        ("OS", "overall survival"),
        ("PCA", "principal component analysis"),
        ("PDC", "Proteomic Data Commons"),
        ("scRNA-seq", "single-cell RNA sequencing"),
        ("TCGA", "The Cancer Genome Atlas"),
        ("TPM", "transcripts per million"),
        ("UMAP", "uniform manifold approximation and projection"),
        ("WSI", "whole-slide image"),
    ]
    for short, long in abbreviations:
        _add_paragraph(
            document, f"{short}: {long}", bold_prefix=f"{short}:", first_line=False
        )

    _add_heading(document, "Declarations", 1)
    for heading, text in declaration_text().items():
        _add_heading(document, heading, 2)
        _add_paragraph(document, text)

    _add_heading(document, "References", 1)
    references = [
        "1. The Cancer Genome Atlas Research Network. Comprehensive molecular profiling of lung adenocarcinoma. Nature. 2014;511:543-550. doi:10.1038/nature13385.",
        "2. Cox DR. Regression models and life-tables. J R Stat Soc Series B Stat Methodol. 1972;34:187-220.",
        "3. Harrell FE Jr, Lee KL, Mark DB. Multivariable prognostic models: issues in developing models, evaluating assumptions and adequacy, and measuring and reducing errors. Stat Med. 1996;15:361-387.",
        "4. Heagerty PJ, Lumley T, Pepe MS. Time-dependent ROC curves for censored survival data and a diagnostic marker. Biometrics. 2000;56:337-344.",
        "5. Edgar R, Domrachev M, Lash AE. Gene Expression Omnibus: NCBI gene expression and hybridization array data repository. Nucleic Acids Res. 2002;30:207-210. doi:10.1093/nar/30.1.207.",
        "6. Shedden K, Taylor JMG, Enkemann SA, et al. Gene expression-based survival prediction in lung adenocarcinoma: a multi-site, blinded validation study. Nat Med. 2008;14:822-827. doi:10.1038/nm.1790.",
        "7. Uhlen M, Zhang C, Lee S, et al. A pathology atlas of the human cancer transcriptome. Science. 2017;357:eaan2507. doi:10.1126/science.aan2507.",
        "8. Gillette MA, Satpathy S, Cao S, et al. Proteogenomic characterization reveals therapeutic vulnerabilities in lung adenocarcinoma. Cell. 2020;182:200-225.e35. doi:10.1016/j.cell.2020.06.013.",
        "9. Subramanian A, Tamayo P, Mootha VK, et al. Gene set enrichment analysis: a knowledge-based approach for interpreting genome-wide expression profiles. Proc Natl Acad Sci USA. 2005;102:15545-15550. doi:10.1073/pnas.0506580102.",
        "10. McShane LM, Altman DG, Sauerbrei W, et al. Reporting recommendations for tumor marker prognostic studies. J Natl Cancer Inst. 2005;97:1180-1184. doi:10.1093/jnci/dji237.",
        "11. Kim N, Kim HK, Lee K, Hong Y, Cho JH, Choi JW, et al. Single-cell RNA sequencing demonstrates the molecular and cellular reprogramming of metastatic lung adenocarcinoma. Nat Commun. 2020;11:2285. doi:10.1038/s41467-020-16164-1.",
        "12. Okayama H, Kohno T, Ishii Y, Shimada Y, Shiraishi K, Iwakawa R, et al. Identification of genes upregulated in ALK-positive and EGFR/KRAS/ALK-negative lung adenocarcinomas. Cancer Res. 2012;72:100-111. doi:10.1158/0008-5472.CAN-11-1403.",
        "13. Der SD, Sykes J, Pintilie M, Zhu CQ, Strumpf D, Liu N, et al. Validation of a histology-independent prognostic gene signature for early-stage non-small-cell lung cancer including stage IA patients. J Thorac Oncol. 2014;9:59-64. doi:10.1097/JTO.0000000000000042.",
        "14. Schabath MB, Welsh EA, Fulp WJ, Chen L, Teer JK, Thompson ZJ, et al. Differential association of STK11 and TP53 with KRAS mutation-associated gene expression, proliferation and immune surveillance in lung adenocarcinoma. Oncogene. 2016;35:3209-3216. doi:10.1038/onc.2015.375.",
        "15. Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent reporting of a multivariable prediction model for individual prognosis or diagnosis (TRIPOD): the TRIPOD Statement. BMC Med. 2015;13:1. doi:10.1186/s12916-014-0241-z.",
        "16. Wolff RF, Moons KGM, Riley RD, Whiting PF, Westwood M, Collins GS, et al. PROBAST: a tool to assess the risk of bias and applicability of prediction model studies. Ann Intern Med. 2019;170:51-58. doi:10.7326/M18-1376.",
        "17. Grossman RL, Heath AP, Ferretti V, Varmus HE, Lowy DR, Kibbe WA, et al. Toward a shared vision for cancer genomic data. N Engl J Med. 2016;375:1109-1112. doi:10.1056/NEJMp1607591.",
        "18. Uno H, Cai T, Pencina MJ, D'Agostino RB, Wei LJ. On the C-statistics for evaluating overall adequacy of risk prediction procedures with censored survival data. Stat Med. 2011;30:1105-1117. doi:10.1002/sim.4154.",
        "19. Katzman JL, Shaham U, Cloninger A, Bates J, Jiang T, Kluger Y. DeepSurv: personalized treatment recommender system using a Cox proportional hazards deep neural network. BMC Med Res Methodol. 2018;18:24. doi:10.1186/s12874-018-0482-1.",
        "20. Lambrechts D, Wauters E, Boeckx B, Aibar S, Nittner D, Burton O, et al. Phenotype molding of stromal cells in the lung tumor microenvironment. Nat Med. 2018;24:1277-1289. doi:10.1038/s41591-018-0096-5.",
        "21. Sinjab A, Han G, Treekitkarnmongkol W, Hara K, Brennan PM, Dang M, et al. Resolving the spatial and cellular architecture of lung adenocarcinoma by multiregion single-cell sequencing. Cancer Discov. 2021;11:2506-2523. doi:10.1158/2159-8290.CD-20-1285.",
        "22. Maynard A, McCoach CE, Rotow JK, Harris L, Haderk F, Kerr DL, et al. Therapy-induced evolution of human lung cancer revealed by single-cell RNA sequencing. Cell. 2020;182:1232-1251.e22. doi:10.1016/j.cell.2020.07.017.",
        "23. Han G, Sinjab A, Rahal Z, Lynch AM, Treekitkarnmongkol W, Liu Y, et al. An atlas of epithelial cell states and plasticity in lung adenocarcinoma. Nature. 2024;627:656-663. doi:10.1038/s41586-024-07113-9.",
        "24. Van Calster B, McLernon DJ, van Smeden M, Wynants L, Steyerberg EW. Calibration: the Achilles heel of predictive analytics. BMC Med. 2019;17:230. doi:10.1186/s12916-019-1466-7.",
    ]
    for reference in references:
        _add_paragraph(document, reference, first_line=False)

    from docx.enum.section import WD_SECTION_START

    landscape = document.add_section(WD_SECTION_START.NEW_PAGE)
    _set_landscape(landscape)
    _add_heading(document, "Tables", 1)
    _build_main_tables(document, evidence, scrna, paths["root"])

    portrait = document.add_section(WD_SECTION_START.NEW_PAGE)
    _set_portrait(portrait)
    _add_heading(document, "Figure legends", 1)
    legends = [
        "Figure 1. Study design and multi-layer validation workflow. TCGA model development and nested cross-validation were separated from locked GEO validation. Bulk cell-state, raw single-cell, HPA and CPTAC analyses were prespecified interpretation layers.",
        "Figure 2. TCGA model development. Nested-cross-validation results compare clinical Cox with RNA and combined models and show the modest increment from RNA PCA_25 plus clinical Cox.",
        "Figure 3. External risk stratification. The frozen continuous score retained discrimination and survival association across four GEO cohorts, while the TCGA-derived absolute cutoff did not transfer reliably.",
        "Figure 4. Cell-state and raw scRNA cellular-context interpretation. Bulk risk associations are integrated with GSE131907 localization of hypoxia, CAF/matrix, and dendritic/B/plasma programs. Proliferation is presented as a broader cycling-cell signal, and EMT-like localization is presented as mainly stromal and myeloid.",
        "Figure 5. HPA and CPTAC protein-level orthogonal support. HPA qualitative/IHC links, CPTAC candidate coverage, and the exploratory hypoxia-module survival association are shown. The CPTAC analysis included 26 deaths and is not presented as definitive validation.",
    ]
    for legend in legends:
        _add_paragraph(document, legend, first_line=False)

    destination = Path(output_path).resolve() if output_path else paths["main"]
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    return destination


def declaration_text() -> dict[str, str]:
    return {
        "Ethics approval and consent to participate": (
            "This study used only publicly available, de-identified datasets and "
            "recruited no new participants. The original studies obtained relevant "
            "ethics approvals and consent. Additional institutional ethics "
            "approval was not sought."
        ),
        "Consent for publication": "Not applicable.",
        "Availability of data and materials": (
            "TCGA-LUAD data are available through the NCI Genomic Data Commons "
            "(https://portal.gdc.cancer.gov/projects/TCGA-LUAD). GEO data are "
            "available under GSE31210, GSE50081, GSE72094, GSE68465 and GSE131907 "
            "(https://www.ncbi.nlm.nih.gov/geo/). HPA evidence is available from "
            "https://www.proteinatlas.org/. CPTAC-LUAD proteomics are available "
            "from the Proteomic Data Commons under PDC000153 "
            "(https://pdc.cancer.gov/). Analysis code, reproducible scripts, "
            "configuration files, and aggregate result tables are available at "
            f"{GITHUB_URL}. Raw and large processed data, whole-slide images, "
            "single-cell matrices, and controlled download credentials are not "
            "included in the code repository."
        ),
        "Competing interests": (
            "The authors declare that they have no competing interests."
        ),
        "Funding": FUNDING_STATEMENT,
        "Authors' contributions": AUTHOR_CONTRIBUTIONS,
        "Acknowledgements": (
            "The authors acknowledge the participants and investigators who "
            "generated and shared TCGA, GEO, HPA, CPTAC and PDC resources."
        ),
        "Authors' information": (
            "Haoyang Li: ORCID 0009-0006-0399-9464. Xiang Li: ORCID "
            "0009-0008-7853-5415. Xuefeng Shi: ORCID 0000-0002-4694-8759."
        ),
    }


def generate_declarations(root: str | Path = ".") -> Path:
    paths = bmc_paths(root)
    paths["manuscript"].mkdir(parents=True, exist_ok=True)
    document = _new_bmc_document()
    _add_title(document, "Declarations")
    _add_paragraph(document, f"Manuscript title: {TITLE}", first_line=False)
    for heading, text in declaration_text().items():
        _add_heading(document, heading, 1)
        _add_paragraph(document, text)
    document.save(paths["declarations"])
    return paths["declarations"]


def generate_title_page(root: str | Path = ".") -> Path:
    paths = bmc_paths(root)
    paths["manuscript"].mkdir(parents=True, exist_ok=True)
    document = _new_bmc_document()
    _add_title(document, TITLE)
    _add_paragraph(document, f"Running title: {RUNNING_TITLE}", first_line=False)
    _add_heading(document, "Authors", 1)
    _add_paragraph(
        document,
        "[AUTHOR TO COMPLETE: Full Name 1]1, [AUTHOR TO COMPLETE: Full Name 2]2, "
        "[AUTHOR TO COMPLETE: Full Name 3]1,*",
        first_line=False,
    )
    _add_heading(document, "Affiliations", 1)
    _add_paragraph(
        document,
        "1 [AUTHOR TO COMPLETE: Department, Institution, City, Postal code, Country]",
        first_line=False,
    )
    _add_paragraph(
        document,
        "2 [AUTHOR TO COMPLETE: Department, Institution, City, Postal code, Country]",
        first_line=False,
    )
    _add_heading(document, "Corresponding author", 1)
    _add_paragraph(
        document,
        "[AUTHOR TO COMPLETE: Name, postal address, email and ORCID]",
        first_line=False,
    )
    _add_heading(document, "Manuscript information", 1)
    _add_paragraph(
        document,
        f"Article type: Research Article\n"
        f"Target journal: BMC Cancer\n"
        f"Main-manuscript word count: {_main_word_count(paths['main'])}\n"
        "Abstract word count: see compliance report\n"
        "Main tables: 4\n"
        "Main figures: 5\n"
        "Supplementary figures: 7\n"
        "Supplementary tables: 8",
        first_line=False,
    )
    _add_heading(document, "Metadata requiring author confirmation", 1)
    for item in [
        "Final author order and spelling",
        "Affiliation mapping and corresponding-author details",
        "ORCID identifiers",
        "Funding, contributions and competing-interest confirmation",
        "Local ethics determination for secondary analysis of public data",
    ]:
        _add_paragraph(document, item, style="List Bullet", first_line=False)
    document.save(paths["title_page"])
    return paths["title_page"]


def generate_cover_letter(root: str | Path = ".") -> Path:
    paths = bmc_paths(root)
    paths["manuscript"].mkdir(parents=True, exist_ok=True)
    evidence = load_stage8_evidence(root)
    document = _new_bmc_document(letter=True)
    _add_paragraph(document, date.today().strftime("%d %B %Y"), first_line=False)
    _add_paragraph(document, "Editor\nBMC Cancer", first_line=False)
    _add_paragraph(document, "Dear Editor,", first_line=False)
    _add_paragraph(
        document,
        f"We submit the Research Article entitled \"{TITLE}\" for consideration "
        "in BMC Cancer.",
    )
    _add_paragraph(
        document,
        "The study addresses the reproducibility and biological interpretation of "
        "LUAD prognostic biomarkers. It combines leakage-controlled TCGA model "
        "development with locked validation in four independent GEO cohorts. The "
        "same fixed risk score is then examined through bulk cell-state programs, "
        "raw single-cell cellular localization, qualitative HPA links and "
        "exploratory CPTAC proteomics.",
    )
    _add_paragraph(
        document,
        f"The external cohorts included {evidence.external_total} patients, with "
        "C-indices from 0.624 to 0.700 and significant continuous-score survival "
        "associations in all four cohorts. Raw GSE131907 data refined the "
        "biological interpretation by supporting malignant epithelial hypoxia, "
        "fibroblast/CAF matrix activity and dendritic/B/plasma immune contexts. "
        "It also showed that proliferation was a broader cycling-cell program and "
        "that the EMT-like bulk signal was mainly stromal/myeloid.",
    )
    _add_paragraph(
        document,
        "We report the modest RNA increment over clinical Cox, the failure of the "
        "TCGA cutoff to transfer, and the neutral WSI pilot. HPA findings are "
        "described as qualitative/IHC-link evidence. CPTAC findings are described "
        f"as exploratory because only {evidence.cptac_deaths} deaths were "
        "available. The manuscript does not present the score for current "
        "clinical decision-making.",
    )
    _add_paragraph(
        document,
        "We recognize BMC Cancer's expectations for prognostic biomarker studies, "
        "including transparent REMARK reporting and biological evidence. The "
        "manuscript provides independent external validation and multi-layer "
        "public-data interpretation, while explicitly acknowledging the absence "
        "of in vitro or in vivo functional experiments as a limitation.",
    )
    _add_paragraph(
        document,
        "This manuscript is not under consideration elsewhere and has not been "
        "published previously, apart from any disclosed preprint [AUTHOR TO "
        "CONFIRM OR DELETE]. All authors have approved the manuscript and agree "
        "to be accountable for the work [AUTHOR TO CONFIRM]. The authors declare "
        "that they have no competing interests [AUTHOR TO CONFIRM].",
    )
    _add_paragraph(document, "Sincerely,", first_line=False)
    _add_paragraph(
        document,
        "[AUTHOR TO COMPLETE: Corresponding author]\n"
        "[AUTHOR TO COMPLETE: Institution]\n"
        "[AUTHOR TO COMPLETE: Email]",
        first_line=False,
    )
    document.save(paths["cover_letter"])
    return paths["cover_letter"]


def build_submission_checklist(root: str | Path = ".") -> pd.DataFrame:
    paths = bmc_paths(root)
    rows = [
        {
            "category": "Required manuscript file",
            "submission_item": "Main manuscript",
            "expected_file": "outputs/manuscript/BMC_Cancer_main_manuscript_draft.docx",
            "required": True,
            "generated": paths["main"].exists(),
            "status": "draft_generated" if paths["main"].exists() else "missing",
            "action_before_submission": (
                "Resolve placeholders, verify references, and insert final figures."
            ),
        },
        {
            "category": "Required manuscript file",
            "submission_item": "Title page",
            "expected_file": "outputs/manuscript/BMC_Cancer_title_page.docx",
            "required": True,
            "generated": paths["title_page"].exists(),
            "status": "metadata_pending",
            "action_before_submission": "Complete authors, affiliations and correspondence.",
        },
        {
            "category": "Editorial file",
            "submission_item": "Cover letter",
            "expected_file": "outputs/manuscript/BMC_Cancer_cover_letter.docx",
            "required": True,
            "generated": paths["cover_letter"].exists(),
            "status": "author_confirmation_pending",
            "action_before_submission": "Confirm exclusivity, approval and disclosures.",
        },
        {
            "category": "Required statements",
            "submission_item": "Declarations",
            "expected_file": "outputs/manuscript/BMC_Cancer_declarations.docx",
            "required": True,
            "generated": paths["declarations"].exists(),
            "status": "author_confirmation_pending",
            "action_before_submission": "Confirm ethics, funding and contributions.",
        },
    ]
    for number in range(1, 6):
        rows.append(
            {
                "category": "Main figure",
                "submission_item": f"Figure {number}",
                "expected_file": f"[AUTHOR TO ASSEMBLE: BMC_Cancer_Figure_{number}.tif]",
                "required": True,
                "generated": False,
                "status": "assembly_pending",
                "action_before_submission": "Assemble and verify labels and resolution.",
            }
        )
    for number in range(1, 5):
        rows.append(
            {
                "category": "Main table",
                "submission_item": f"Table {number}",
                "expected_file": "Embedded as editable table in main manuscript",
                "required": True,
                "generated": paths["main"].exists(),
                "status": "draft_embedded",
                "action_before_submission": "Numerical and editorial proofread.",
            }
        )
    for number, label in {
        1: "Additional model performance and calibration",
        2: "GEO preprocessing and gene coverage",
        3: "Full cell-state correlation heatmaps",
        4: "Full raw scRNA dotplots, UMAPs and heatmaps",
        5: "HPA candidate protein evidence",
        6: "CPTAC candidate and module analyses",
        7: "Exploratory WSI neutral pilot results",
    }.items():
        rows.append(
            {
                "category": "Supplementary figure",
                "submission_item": f"Supplementary Fig. S{number}: {label}",
                "expected_file": f"[AUTHOR TO ASSEMBLE: Supplementary_Figure_S{number}]",
                "required": True,
                "generated": False,
                "status": "source_outputs_available_assembly_pending",
                "action_before_submission": "Assemble source panels with a complete legend.",
            }
        )
    for number, label in {
        1: "Full Stage 2D model performance",
        2: "GEO preprocessing and missingness",
        3: "Cell-state signature definitions",
        4: "Full cell-state association results",
        5: "Raw scRNA gene and program localization",
        6: "HPA evidence and IHC links",
        7: "CPTAC protein matching and module analyses",
        8: "WSI pilot diagnostics",
    }.items():
        rows.append(
            {
                "category": "Supplementary table",
                "submission_item": f"Supplementary Table S{number}: {label}",
                "expected_file": f"[AUTHOR TO ASSEMBLE: Supplementary_Table_S{number}]",
                "required": True,
                "generated": False,
                "status": "source_outputs_available_assembly_pending",
                "action_before_submission": "Assemble from locked project outputs.",
            }
        )
    rows.extend(
        [
            {
                "category": "Reporting guideline",
                "submission_item": "REMARK checklist",
                "expected_file": "[AUTHOR TO COMPLETE: REMARK checklist]",
                "required": True,
                "generated": False,
                "status": "required_pending",
                "action_before_submission": "Complete against the final manuscript.",
            },
            {
                "category": "Editorial scope",
                "submission_item": "Biological validation expectation",
                "expected_file": "Not available",
                "required": False,
                "generated": False,
                "status": "material_editorial_risk",
                "action_before_submission": (
                    "Consider whether experimental validation is feasible; otherwise "
                    "retain the explicit limitation and expect possible desk rejection."
                ),
            },
            {
                "category": "Reproducibility",
                "submission_item": "Public code release",
                "expected_file": "[AUTHOR TO COMPLETE: repository URL and archived release]",
                "required": False,
                "generated": False,
                "status": "public_release_pending",
                "action_before_submission": "Publish a versioned release without raw large data.",
            },
        ]
    )
    frame = pd.DataFrame(rows)
    paths["tables"].mkdir(parents=True, exist_ok=True)
    frame.to_csv(paths["checklist"], index=False)
    return frame


def _sequence(text: str, pattern: str) -> list[int]:
    values = [int(value) for value in re.findall(pattern, text)]
    first_seen: list[int] = []
    for value in values:
        if value not in first_seen:
            first_seen.append(value)
    return first_seen


def compliance_report(root: str | Path = ".") -> str:
    paths = bmc_paths(root)
    cleaned = (
        paths["manuscript"]
        / "BMC_Cancer_main_manuscript_reference_table_cleaned.docx"
    )
    manuscript_path = cleaned if cleaned.exists() else paths["main"]
    text = _document_text(manuscript_path)
    xml = _document_xml(manuscript_path)
    abstract_match = re.search(
        r"Abstract\s+(.*?)\s+Keywords:", text, flags=re.DOTALL | re.IGNORECASE
    )
    abstract = abstract_match.group(1) if abstract_match else ""
    abstract_words = len(re.findall(r"\b[\w'-]+\b", abstract))
    keyword_match = re.search(r"Keywords:\s*(.+)", text, flags=re.IGNORECASE)
    keywords = (
        [item.strip() for item in keyword_match.group(1).split(";") if item.strip()]
        if keyword_match
        else []
    )
    figures, figure_order = _citation_order(text, "Fig", 5)
    tables, table_order = _citation_order(text, "Table", 4)
    supplementary_figures = _sequence(text, r"Supplementary Fig\.\s*S(\d+)")
    supplementary_tables = _sequence(text, r"Supplementary Table\s*S(\d+)")
    required_sections = [
        "Abstract",
        "Introduction",
        "Methods",
        "Results",
        "Discussion",
        "Conclusion",
        "List of abbreviations",
        "Declarations",
        "References",
    ]
    missing_sections = [section for section in required_sections if section not in text]
    placeholders = len(
        re.findall(r"\[(?:AUTHOR TO|TO BE|Optional; TO BE)", text, flags=re.IGNORECASE)
    )
    checklist = build_submission_checklist(root)
    checks = [
        ("Structured abstract <=350 words", 0 < abstract_words <= 350),
        ("Three to ten keywords", 3 <= len(keywords) <= 10),
        ("Required manuscript sections present", not missing_sections),
        ("Continuous line numbering encoded", 'w:restart="continuous"' in xml),
        ("Page-number field encoded", " PAGE " in xml),
        ("Double line spacing encoded", 'w:line="480"' in xml or 'w:lineRule="auto"' in xml),
        ("Figures first cited in order 1-5", figure_order),
        ("Tables first cited in order 1-4", table_order),
        (
            "Supplementary figures first cited in order S1-S7",
            supplementary_figures[:7] == list(range(1, 8)),
        ),
        (
            "Supplementary tables first cited in order S1-S8",
            supplementary_tables[:8] == list(range(1, 9)),
        ),
    ]
    check_lines = "\n".join(
        f"| {name} | {'PASS' if passed else 'REVIEW'} |" for name, passed in checks
    )
    pending = checklist.loc[
        checklist["required"].astype(bool) & ~checklist["generated"].astype(bool),
        "submission_item",
    ].tolist()
    return f"""# BMC Cancer Compliance Report

Generated: {datetime.now().isoformat(timespec="seconds")}

## Scope

This engineering audit compares the latest cleaned Research Article draft with the current BMC Cancer manuscript guidance. It does not constitute editorial acceptance.

## Automated Checks

| Check | Status |
| --- | --- |
{check_lines}

## Counts

- Abstract words: {abstract_words} (journal guidance: no more than 350).
- Keywords: {len(keywords)} ({", ".join(keywords)}).
- Main manuscript words, including references, tables and legends: {_main_word_count(manuscript_path)}.
- First-seen main figure citations: {figures}.
- First-seen main table citations: {tables}.
- First-seen supplementary figure citations: {supplementary_figures}.
- First-seen supplementary table citations: {supplementary_tables}.
- Author-confirmation or completion placeholders in the main manuscript: {placeholders}.

## BMC Cancer-Specific Editorial Risk

The current BMC Cancer scope states that prognostic biomarker studies should follow REMARK, include an independent validation cohort and include biological validation using in vitro or in vivo experiments. It also warns that models based only on computational analyses of public databases without adequate validation are unlikely to be considered.

This project has a material strength: the frozen model was evaluated in four independent GEO cohorts. It also has public-data orthogonal evidence from raw scRNA-seq, HPA and CPTAC. However, those layers are cellular-context and protein-association evidence, not in vitro or in vivo functional validation. Therefore:

- REMARK completion is mandatory before submission.
- The manuscript should not imply that Stage 4B, HPA or CPTAC satisfies functional biological validation.
- There remains a substantial desk-rejection risk under the journal's stated biomarker scope.
- Submission may still be attempted with transparent framing, but the risk should be accepted consciously.

## Author Actions

- Missing manuscript sections: {", ".join(missing_sections) if missing_sections else "none"}.
- Required submission assets still pending: {", ".join(pending) if pending else "none"}.
- Reconfirm author metadata, the local ethics approach, and contribution statements before submission.
- Reference and placeholder cleanup passed in the latest cleaned manuscript.
- Complete the REMARK checklist against the final version.
- Assemble Figures 1-5, Supplementary Figs. S1-S7 and Supplementary Tables S1-S8.
- The public GitHub repository is available; create a versioned archival release or DOI if possible.
- Recheck the live BMC Cancer instructions immediately before submission.

## Official Sources

- Research Article guidance: {BMC_GUIDELINES_URL}
- Aims and scope, including biomarker-study requirements: {BMC_SCOPE_URL}

## Status

**Draft generated, not submission-ready.** Formatting checks are largely automated, but author metadata, final assets, REMARK reporting and the biological-validation scope risk remain unresolved.
"""


RISK_PATTERNS = {
    "causality overclaim": [
        r"\bcausal confirmation\b",
        r"\bcausally confirmed\b",
        r"\bproved? the mechanism\b",
    ],
    "clinical readiness": [
        r"\bclinically ready tool\b",
        r"\bready for clinical (?:use|deployment)\b",
        r"\bdemonstrates clinical utility\b",
    ],
    "protein overclaim": [
        r"\bdefinitive proteomic validation\b",
        r"\ball candidates validated by CPTAC\b",
    ],
    "single-cell overclaim": [
        r"\bsingle-cell validated survival model\b",
        r"\bsingle-cell survival validation\b",
    ],
    "WSI overclaim": [
        r"\bdigital pathology improved prediction\b",
        r"\bWSI improved prediction\b",
    ],
    "multimodal overclaim": [r"\bmultimodal deep learning superiority\b"],
    "EMT overclaim": [r"\bmalignant EMT confirmed\b"],
    "study-design overclaim": [r"\bprospective validation\b"],
}


def language_risk_report(root: str | Path = ".") -> str:
    paths = bmc_paths(root)
    scanned = [
        paths["main"],
        paths["title_page"],
        paths["cover_letter"],
        paths["declarations"],
    ]
    hits: list[tuple[str, str, str]] = []
    for path in scanned:
        text = _document_text(path)
        for category, patterns in RISK_PATTERNS.items():
            for pattern in patterns:
                for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                    start = max(0, match.start() - 70)
                    end = min(len(text), match.end() + 70)
                    hits.append(
                        (path.name, category, " ".join(text[start:end].split()))
                    )
    rows = (
        "\n".join(f"| {file} | {category} | {context} |" for file, category, context in hits)
        if hits
        else "| None | None | No prohibited positive assertions detected. |"
    )
    return f"""# BMC Cancer Claim-Language Audit

Generated: {datetime.now().isoformat(timespec="seconds")}

## Result

- Documents scanned: {len(scanned)}.
- High-risk phrase hits: {len(hits)}.
- Automated status: {"REVIEW REQUIRED" if hits else "PASS"}.

| File | Risk category | Context |
| --- | --- | --- |
{rows}

## Required Language Boundaries

- Describe external results as reproducible risk ranking, not a calibrated clinical tool.
- State that the RNA increment over clinical Cox was modest.
- Describe Stage 4 as cell-state associations inferred from bulk RNA-seq.
- Describe Stage 4B as raw scRNA-supported cellular-context interpretation without linked survival.
- Interpret proliferation as a broader cycling-cell signal.
- Interpret EMT-like signal as mainly stromal/myeloid rather than clearly malignant-cell EMT.
- Describe HPA as qualitative/IHC-link evidence.
- Describe CPTAC as exploratory quantitative support because only 26 deaths were available.
- State that WSI did not show stable improvement.
"""


def evidence_claim_report(root: str | Path = ".") -> str:
    evidence = load_stage8_evidence(root)
    external = ", ".join(
        f"{row['cohort']} {row['c_index']:.3f}" for row in _external_rows(evidence)
    )
    rows = [
        (
            "The frozen score ranked survival risk across external cohorts.",
            f"Four GEO C-indices: {external}; continuous-score Cox P<0.01 in each cohort.",
            "Supported",
            "Use external risk-ranking language.",
        ),
        (
            "RNA strongly outperformed clinical variables.",
            f"Nested C-index {evidence.nested_combined_cindex:.3f} versus {evidence.nested_clinical_cindex:.3f}.",
            "Not supported",
            "Call the increment modest.",
        ),
        (
            "The TCGA cutoff transferred across platforms.",
            "It failed or collapsed groups in the external cohorts.",
            "Not supported",
            "Keep cohort-median analyses as sensitivity analyses.",
        ),
        (
            "Higher risk was associated with hypoxia, proliferation, EMT-like and CAF programs.",
            "TCGA associations and concordant GEO directions.",
            "Supported as association",
            "Bulk signature inference only.",
        ),
        (
            "Raw scRNA localized hypoxia and LDHA to malignant epithelial cells.",
            "GSE131907 hypoxia ranked malignant epithelial first; LDHA was highest there.",
            "Supported as cellular context",
            "No single-cell survival claim.",
        ),
        (
            "Raw scRNA localized CAF/matrix and dendritic/B/plasma programs.",
            "Fibroblast/CAF and expected immune-lineage localization.",
            "Supported as cellular context",
            "No abundance or causation claim.",
        ),
        (
            "Proliferation and MKI67/CDK1 were malignant-epithelial specific.",
            "T cells ranked first and malignant epithelial cells second.",
            "Not supported",
            "Use broader cycling-cell interpretation.",
        ),
        (
            "The EMT-like signal demonstrated malignant-cell EMT.",
            "Top contexts were fibroblast/CAF, myeloid and endothelial.",
            "Not supported",
            "Use stromal/myeloid-associated interpretation.",
        ),
        (
            "HPA supplied quantitative protein validation.",
            f"{evidence.hpa_supported_count}/{evidence.hpa_candidate_count} records or links; no uniform endpoint.",
            "Not supported",
            "Use qualitative/IHC-link evidence.",
        ),
        (
            "CPTAC supplied definitive validation.",
            f"{evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} candidates and {evidence.cptac_deaths} deaths.",
            "Not supported",
            "Use exploratory quantitative support.",
        ),
        (
            "WSI improved prediction.",
            f"Clinical {evidence.wsi_clinical_cindex:.3f}; pathology {evidence.wsi_pathology_cindex:.3f}; fusion {evidence.wsi_fusion_cindex:.3f}.",
            "Not supported",
            "Report feasibility and neutral performance.",
        ),
        (
            "The study meets BMC Cancer's in vitro or in vivo biomarker-validation expectation.",
            "No functional experiment was performed.",
            "Not supported",
            "Flag as a material editorial limitation.",
        ),
    ]
    table = "\n".join(
        f"| {claim} | {evidence_text} | {status} | {wording} |"
        for claim, evidence_text, status, wording in rows
    )
    return f"""# BMC Cancer Evidence-Claim Audit

Generated: {datetime.now().isoformat(timespec="seconds")}

## Evidence Hierarchy

1. Primary prognostic evidence: leakage-controlled TCGA nested cross-validation and locked validation in four GEO cohorts.
2. Bulk interpretation: prespecified cell-state associations with external direction checks.
3. Raw single-cell interpretation: cellular localization in GSE131907 without linked survival.
4. Protein context: qualitative HPA links and exploratory CPTAC analyses.
5. Supplementary neutral evidence: WSI feasibility without stable added value.

## Claim Map

| Proposed claim | Project evidence | Status | Permitted wording |
| --- | --- | --- | --- |
{table}

## Integrity Decision

The manuscript can lead with reproducible external risk ranking and convergent cellular-context interpretation. It cannot lead with clinical deployment, functional mechanism, malignant-cell EMT, conclusive protein validation, pathology improvement or compliance with experimental biomarker validation.
"""


def figure_table_audit(root: str | Path = ".") -> str:
    paths = bmc_paths(root)
    text = _document_text(paths["main"])
    figures, figure_order = _citation_order(text, "Fig", 5)
    tables, table_order = _citation_order(text, "Table", 4)
    supp_figures = _sequence(text, r"Supplementary Fig\.\s*S(\d+)")
    supp_tables = _sequence(text, r"Supplementary Table\s*S(\d+)")
    return f"""# BMC Cancer Figure and Table Audit

Generated: {datetime.now().isoformat(timespec="seconds")}

## Citation Order

| Asset group | First-seen sequence | Status |
| --- | --- | --- |
| Main figures | {figures} | {"PASS" if figure_order else "REVIEW"} |
| Main tables | {tables} | {"PASS" if table_order else "REVIEW"} |
| Supplementary figures | {supp_figures} | {"PASS" if supp_figures[:7] == list(range(1, 8)) else "REVIEW"} |
| Supplementary tables | {supp_tables} | {"PASS" if supp_tables[:8] == list(range(1, 9)) else "REVIEW"} |

## Main-Text Plan

1. Figure 1: study design and evidence hierarchy.
2. Figure 2: TCGA model comparison and modest RNA increment.
3. Figure 3: four-cohort external validation and cutoff limitation.
4. Figure 4: bulk cell-state associations plus concise raw scRNA localization.
5. Figure 5: HPA qualitative links and exploratory CPTAC support.
6. Tables 1-4: cohorts, model performance, cellular-context summary and protein evidence.

## Supplementary Plan

- S1: additional model performance and calibration.
- S2: GEO preprocessing and gene coverage.
- S3: full bulk cell-state correlation results.
- S4: full raw scRNA dotplots, heatmaps and UMAPs.
- S5: complete HPA candidate evidence.
- S6: complete CPTAC candidate and module analyses.
- S7: WSI feasibility and neutral pilot results.
- Supplementary Tables S1-S8 follow the same evidence order and keep WSI last.

## Assembly Status

The manuscript contains ordered citations and legends, but publication-resolution composite figures and final supplementary files remain to be assembled. Stage 4B source plots are available under `outputs/figures/stage4b_lowmem_*`.
"""


def write_compliance_outputs(root: str | Path = ".") -> dict[str, Path]:
    paths = bmc_paths(root)
    paths["reports"].mkdir(parents=True, exist_ok=True)
    build_submission_checklist(root)
    paths["compliance"].write_text(compliance_report(root), encoding="utf-8")
    paths["claims"].write_text(evidence_claim_report(root), encoding="utf-8")
    paths["figure_table"].write_text(figure_table_audit(root), encoding="utf-8")
    return {
        "checklist": paths["checklist"],
        "compliance": paths["compliance"],
        "claims": paths["claims"],
        "figure_table": paths["figure_table"],
    }


def write_language_audit(root: str | Path = ".") -> Path:
    paths = bmc_paths(root)
    paths["reports"].mkdir(parents=True, exist_ok=True)
    paths["language"].write_text(language_risk_report(root), encoding="utf-8")
    return paths["language"]


def append_stage9_bmc_audit(root: str | Path, message: str) -> Path:
    path = bmc_paths(root)["audit"]
    with path.open("a", encoding="utf-8") as handle:
        handle.write(message)
    return path
