"""Stage 8 manuscript, figure/table planning, and evidence-audit utilities."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
import re
import shutil

import pandas as pd


class Stage8Error(RuntimeError):
    """Raised when required evidence for Stage 8 is unavailable."""


@dataclass(frozen=True)
class Stage8Evidence:
    tcga_patients: int
    tcga_deaths: int
    tcga_censored: int
    nested_combined_cindex: float
    nested_combined_sd: float
    nested_clinical_cindex: float
    nested_clinical_sd: float
    nested_auc_1y: float
    nested_auc_3y: float
    nested_auc_5y: float
    external_total: int
    external_performance: pd.DataFrame
    external_cox: pd.DataFrame
    cell_state_correlations: pd.DataFrame
    cell_state_consistency: pd.DataFrame
    adjusted_cell_states: list[str]
    signature_count: int
    hpa_candidate_count: int
    hpa_supported_count: int
    cptac_samples: int
    cptac_proteins: int
    cptac_usable_os: int
    cptac_deaths: int
    cptac_candidates_matched: int
    cptac_candidate_total: int
    cptac_supported_genes: list[str]
    cptac_hypoxia_hr: float
    cptac_hypoxia_ci_low: float
    cptac_hypoxia_ci_high: float
    cptac_hypoxia_p: float
    cptac_hypoxia_fdr: float
    wsi_slides: int
    wsi_clinical_cindex: float
    wsi_pathology_cindex: float
    wsi_fusion_cindex: float


def stage8_paths(root: str | Path = ".") -> dict[str, Path]:
    project_root = Path(root).resolve()
    return {
        "root": project_root,
        "manuscript_dir": project_root / "outputs" / "manuscript",
        "tables_dir": project_root / "outputs" / "tables",
        "reports_dir": project_root / "outputs" / "reports",
        "figures_dir": project_root / "outputs" / "figures",
        "audit": project_root / "audit_report.md",
    }


def _read_required(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(f"Required Stage 8 evidence file is missing: {path}")
    frame = pd.read_csv(path)
    if frame.empty:
        raise Stage8Error(f"Required Stage 8 evidence table is empty: {path}")
    return frame


def load_stage8_evidence(root: str | Path = ".") -> Stage8Evidence:
    paths = stage8_paths(root)
    tables = paths["tables_dir"]
    feature = _read_required(tables / "stage2c_feature_space_comparison.csv")
    km = _read_required(tables / "stage2c_oof_km_logrank.csv")
    external = _read_required(tables / "stage2d_external_validation_performance.csv")
    external_cox = _read_required(tables / "stage2d_external_validation_cox.csv")
    geo_summary = _read_required(tables / "stage2d_geo_dataset_summary.csv")
    correlations = _read_required(tables / "stage4_tcga_cell_state_risk_correlation.csv")
    consistency = _read_required(tables / "stage4_cell_state_external_consistency.csv")
    adjusted = _read_required(tables / "stage4_tcga_cell_state_multivariable_cox.csv")
    signatures = _read_required(tables / "stage4_cell_state_signature_definitions.csv")
    hpa = _read_required(tables / "stage5_hpa_protein_evidence.csv")
    inventory = _read_required(tables / "stage5b_cptac_data_inventory.csv")
    availability = _read_required(tables / "stage5b_cptac_candidate_availability.csv")
    integrated = _read_required(tables / "stage5b_integrated_hpa_cptac_evidence.csv")
    module_cox = _read_required(tables / "stage5b_cptac_protein_module_survival_cox.csv")
    wsi = _read_required(tables / "stage6a_wsi_pilot_model_performance.csv")
    wsi_low_complexity = _read_required(
        tables / "stage6a_low_complexity_pathology_performance.csv"
    )

    combined = feature.loc[feature["config_id"].eq("rna_pca25_clinical_cox")]
    clinical = feature.loc[feature["config_id"].eq("clinical_cox")]
    if combined.empty or clinical.empty:
        raise Stage8Error("Stage 2C PCA_25 combined or clinical Cox evidence is missing.")
    combined_row = combined.iloc[0]
    clinical_row = clinical.iloc[0]
    tcga_rows = km.loc[km["config_id"].eq("rna_pca25_clinical_cox")]
    if tcga_rows.empty:
        raise Stage8Error("Stage 2C OOF evidence for the frozen configuration is missing.")

    external_full = external.loc[external["analysis_type"].eq("frozen_full_model")].copy()
    external_full = external_full.sort_values("cohort").reset_index(drop=True)
    external_cox_full = external_cox.loc[
        external_cox["analysis_type"].eq("frozen_full_model")
    ].sort_values("cohort").reset_index(drop=True)

    adjusted_risk = adjusted.loc[adjusted["covariate"].eq("risk_score")].copy()
    adjusted_states = adjusted_risk.loc[
        pd.to_numeric(adjusted_risk["p_value"], errors="coerce").lt(0.05),
        "signature_name",
    ].astype(str).tolist()

    protein_inventory = inventory.loc[
        inventory["role"].eq("protein_abundance_matrix")
    ]
    if protein_inventory.empty:
        raise Stage8Error("CPTAC protein inventory is unavailable for Stage 8.")
    protein_inventory_row = protein_inventory.iloc[0]
    clinical_processed = _read_required(
        paths["root"] / "data" / "processed" / "cptac_luad" / "cptac_luad_clinical_processed.csv"
    )
    usable_os = clinical_processed[["os_time_days", "os_event"]].notna().all(axis=1)
    hypoxia = module_cox.loc[
        module_cox["mechanism_layer_score"].eq("hypoxia")
    ]
    if hypoxia.empty:
        raise Stage8Error("CPTAC hypoxia module Cox result is missing.")
    hypoxia_row = hypoxia.iloc[0]
    cptac_supported = integrated.loc[
        integrated["supported_by_CPTAC_quantitative"].fillna(False).astype(bool),
        "gene_symbol",
    ].astype(str).tolist()

    test_wsi = wsi.loc[wsi["split"].eq("test")]
    clinical_wsi = test_wsi.loc[test_wsi["model_name"].eq("clinical_only_cox")]
    fusion_wsi = test_wsi.loc[
        test_wsi["model_name"].str.contains("fusion", case=False, na=False)
    ]
    low_complexity_test = wsi_low_complexity.loc[
        wsi_low_complexity["split"].eq("test")
    ]
    pathology_wsi = low_complexity_test.loc[
        low_complexity_test["model_name"].str.contains("pathology", case=False, na=False)
        & ~low_complexity_test["model_name"].str.contains("clinical", case=False, na=False)
    ]
    if clinical_wsi.empty or pathology_wsi.empty or fusion_wsi.empty:
        raise Stage8Error("Required WSI pilot comparison rows are missing.")

    return Stage8Evidence(
        tcga_patients=int(tcga_rows["patient_count"].max()),
        tcga_deaths=int(tcga_rows["death_count"].max()),
        tcga_censored=int(tcga_rows["patient_count"].max() - tcga_rows["death_count"].max()),
        nested_combined_cindex=float(combined_row["outer_c_index_mean"]),
        nested_combined_sd=float(combined_row["outer_c_index_std"]),
        nested_clinical_cindex=float(clinical_row["outer_c_index_mean"]),
        nested_clinical_sd=float(clinical_row["outer_c_index_std"]),
        nested_auc_1y=float(combined_row["auc_1_year_mean"]),
        nested_auc_3y=float(combined_row["auc_3_year_mean"]),
        nested_auc_5y=float(combined_row["auc_5_year_mean"]),
        external_total=int(geo_summary["patients"].sum()),
        external_performance=external_full,
        external_cox=external_cox_full,
        cell_state_correlations=correlations.sort_values(
            "spearman_rho", ascending=False
        ).reset_index(drop=True),
        cell_state_consistency=consistency,
        adjusted_cell_states=adjusted_states,
        signature_count=int(signatures["signature_name"].nunique()),
        hpa_candidate_count=int(hpa["gene_symbol"].nunique()),
        hpa_supported_count=int(
            hpa["hpa_query_status"].astype(str).str.lower().ne("failed").sum()
        ),
        cptac_samples=int(protein_inventory_row["row_count"]),
        cptac_proteins=int(protein_inventory_row["column_count"] - 1),
        cptac_usable_os=int(usable_os.sum()),
        cptac_deaths=int(clinical_processed.loc[usable_os, "os_event"].sum()),
        cptac_candidates_matched=int(availability["cptac_available"].fillna(False).sum()),
        cptac_candidate_total=int(len(availability)),
        cptac_supported_genes=cptac_supported,
        cptac_hypoxia_hr=float(hypoxia_row["hazard_ratio_per_sd"]),
        cptac_hypoxia_ci_low=float(hypoxia_row["ci95_lower"]),
        cptac_hypoxia_ci_high=float(hypoxia_row["ci95_upper"]),
        cptac_hypoxia_p=float(hypoxia_row["p_value"]),
        cptac_hypoxia_fdr=float(hypoxia_row["fdr_bh"]),
        wsi_slides=100,
        wsi_clinical_cindex=float(clinical_wsi.iloc[0]["c_index"]),
        wsi_pathology_cindex=float(pathology_wsi["c_index"].max()),
        wsi_fusion_cindex=float(fusion_wsi["c_index"].max()),
    )


def _external_lines(evidence: Stage8Evidence) -> str:
    lines = []
    cox = evidence.external_cox.set_index("cohort")
    for row in evidence.external_performance.itertuples(index=False):
        cox_row = cox.loc[row.cohort]
        lines.append(
            f"- {row.cohort}: n={int(row.patient_count)}, deaths={int(row.death_events)}, "
            f"C-index={row.c_index:.3f}, 1/3/5-year AUC="
            f"{row.auc_1_year:.3f}/{row.auc_3_year:.3f}/{row.auc_5_year:.3f}, "
            f"risk-score HR per SD={cox_row['hazard_ratio_per_sd']:.3f}, "
            f"P={cox_row['p_value']:.3g}."
        )
    return "\n".join(lines)


def _top_cell_state_text(evidence: Stage8Evidence, limit: int = 5) -> str:
    top = evidence.cell_state_correlations.head(limit)
    return ", ".join(
        f"{row.signature_name} (rho={row.spearman_rho:.3f})"
        for row in top.itertuples(index=False)
    )


def manuscript_skeleton_markdown(evidence: Stage8Evidence) -> str:
    external_c = ", ".join(
        f"{row.cohort} {row.c_index:.3f}"
        for row in evidence.external_performance.itertuples(index=False)
    )
    return f"""# SC-PROST-LUAD Manuscript Skeleton

## Working Title

**An externally validated transcriptomic-clinical survival model with cell-state and proteomic interpretation in lung adenocarcinoma**

## Alternative Title

**An externally validated transcriptomic-clinical survival model reveals hypoxia, proliferation, EMT-like and stromal programs in lung adenocarcinoma**

## One-Sentence Argument

In lung adenocarcinoma, a frozen TCGA-derived transcriptomic-clinical survival score reproducibly ranked risk across four independent GEO cohorts, while bulk RNA-derived cell-state associations and exploratory CPTAC proteomics linked higher risk to hypoxia and proliferative programs, within the limits of retrospective public data and modest incremental value over clinical variables.

## Abstract Draft

### Background

Transcriptomic prognostic models for lung adenocarcinoma frequently lack external validation and biologically bounded interpretation across independent platforms.

### Methods

We developed a leakage-controlled overall-survival model in {evidence.tcga_patients} TCGA-LUAD patients using clinical covariates and a PCA-based RNA representation. The frozen model was evaluated without refitting in four GEO cohorts comprising {evidence.external_total} patients. We interpreted the risk score using {evidence.signature_count} curated LUAD cell-state signatures, Human Protein Atlas evidence for {evidence.hpa_candidate_count} prespecified candidates, and PDC CPTAC-LUAD proteomics.

### Results

In nested cross-validation, the RNA PCA_25 plus clinical Cox model achieved a C-index of {evidence.nested_combined_cindex:.3f} +/- {evidence.nested_combined_sd:.3f}, compared with {evidence.nested_clinical_cindex:.3f} +/- {evidence.nested_clinical_sd:.3f} for clinical Cox. External C-indices were {external_c}, and the continuous frozen risk score was associated with overall survival in all four cohorts. The TCGA-derived absolute cutoff did not transfer reliably across microarray platforms. Higher risk was associated with hypoxia, proliferation, malignant epithelial, EMT-like and CAF signatures, with directional support for the principal programs across GEO cohorts. HPA provided qualitative or linkable evidence for all {evidence.hpa_supported_count} candidates. In {evidence.cptac_samples} CPTAC primary tumors, {evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} candidate proteins were quantified; LDHA, MKI67 and CDK1 showed nominally significant direction-consistent stage associations, and the hypoxia protein module was associated with survival (HR {evidence.cptac_hypoxia_hr:.2f} per SD, 95% CI {evidence.cptac_hypoxia_ci_low:.2f}-{evidence.cptac_hypoxia_ci_high:.2f}, FDR={evidence.cptac_hypoxia_fdr:.4f}).

### Conclusions

The frozen score showed reproducible external risk ranking and was associated with coherent hypoxia, proliferative, EMT-like and stromal programs. The modest RNA increment in TCGA, platform-dependent cutoff transfer, signature-based cell-state inference and limited CPTAC events support further prospective and assay-harmonized validation before clinical use.

## Introduction

### Paragraph 1: Clinical Need

- Establish the prognostic heterogeneity of LUAD after diagnosis and treatment.
- Explain why clinicopathological variables alone do not fully capture molecular risk.
- Add literature citations during manuscript completion.

### Paragraph 2: Current Modeling Gap

- Summarize transcriptomic survival modeling and common limitations: leakage, weak external validation, platform shift and limited interpretation.
- Distinguish risk ranking from calibrated absolute risk prediction.

### Paragraph 3: Biological Interpretation Gap

- Explain why bulk RNA risk scores require tumor-state and microenvironment interpretation.
- Position curated single-cell-derived signatures as an association layer rather than direct single-cell outcome validation.
- Introduce orthogonal HPA and CPTAC evidence as supporting, not mechanistic, evidence.

### Paragraph 4: Present Study

- State the frozen TCGA-to-GEO validation design.
- State the three evidence layers: external risk ranking, cell-state association and protein-level support.
- State that exploratory WSI analyses were retained as a feasibility/limitation analysis because pathology models did not improve on the clinical reference.

## Results

### 1. Leakage-Controlled TCGA Model Development Identified a Modest RNA Increment

- TCGA cohort: {evidence.tcga_patients} patients, {evidence.tcga_deaths} deaths and {evidence.tcga_censored} censored observations.
- Nested CV: RNA PCA_25 + clinical Cox C-index {evidence.nested_combined_cindex:.3f} +/- {evidence.nested_combined_sd:.3f}; clinical Cox {evidence.nested_clinical_cindex:.3f} +/- {evidence.nested_clinical_sd:.3f}.
- 1/3/5-year nested-CV AUC: {evidence.nested_auc_1y:.3f}/{evidence.nested_auc_3y:.3f}/{evidence.nested_auc_5y:.3f}.
- Interpretation: measurable but small incremental discrimination; DeepSurv overfitting precluded a neural-network superiority claim.

### 2. The Frozen Score Reproduced Risk Ranking Across Four GEO Cohorts

{_external_lines(evidence)}

- The continuous score was significant in univariable Cox analyses in all cohorts.
- The absolute TCGA median cutoff was not transportable; cohort-median KM analyses remain sensitivity analyses.
- Calibration and clinical deployment claims are outside the supported evidence.

### 3. Higher RNA Risk Was Associated with Tumor-State and Stromal Programs

- Strongest positive TCGA associations: {_top_cell_state_text(evidence)}.
- Core externally supported adverse programs: EMT-like, proliferation, hypoxia and CAF.
- Dendritic/B/plasma signatures showed the prespecified low-risk direction.
- M2 macrophage, Treg, exhausted CD8, cytotoxic CD8, M1 macrophage and NK findings remain non-core or unstable.
- Adjusted TCGA associations with P<0.05: {", ".join(evidence.adjusted_cell_states)}.

### 4. HPA and CPTAC Provided Bounded Protein-Level Support

- HPA qualitative/link evidence: {evidence.hpa_supported_count}/{evidence.hpa_candidate_count} candidates.
- CPTAC-LUAD: {evidence.cptac_samples} primary tumors, {evidence.cptac_proteins} quantified protein genes, {evidence.cptac_usable_os} usable OS records and {evidence.cptac_deaths} deaths.
- Candidate matching: {evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total}.
- Statistically supported stage-direction candidates: {", ".join(evidence.cptac_supported_genes)}.
- Hypoxia module survival association: HR {evidence.cptac_hypoxia_hr:.3f}, 95% CI {evidence.cptac_hypoxia_ci_low:.3f}-{evidence.cptac_hypoxia_ci_high:.3f}, P={evidence.cptac_hypoxia_p:.4f}, FDR={evidence.cptac_hypoxia_fdr:.4f}.
- No individual candidate protein remained significant after survival FDR correction.

### 5. Exploratory WSI Modeling Was Technically Feasible but Did Not Improve Prediction

- OpenSlide reading, patch extraction, CUDA ResNet50 feature extraction and MIL training completed for a {evidence.wsi_slides}-slide pilot.
- Test C-index: clinical-only {evidence.wsi_clinical_cindex:.3f}; best pathology-only {evidence.wsi_pathology_cindex:.3f}; best fusion {evidence.wsi_fusion_cindex:.3f}.
- Keep this analysis supplementary and interpret it as a neutral feasibility result.

## Discussion

### Central Finding

The study supports external risk ranking by a frozen transcriptomic-clinical score and connects that ranking to coherent hypoxia, proliferation, EMT-like and stromal associations.

### Interpretation

- External discrimination was consistent, but the RNA increment over clinical Cox in TCGA was modest.
- Continuous-score transport was stronger than absolute-cutoff transport.
- Cell-state findings are compatible with high-risk tumor-state and stromal biology, not evidence of causality.
- CPTAC results prioritize hypoxia and proliferation for follow-up, with the hypoxia module providing the clearest exploratory survival signal.

### Clinical Meaning

- Present the score as a research stratification signal.
- Do not present it as a deployed decision tool.
- Prospective validation, assay harmonization, cutoff prespecification and calibration are required.

### WSI Finding

The WSI pilot demonstrates engineering feasibility but also provides an important negative result: added pathology complexity did not yield stable improvement over the clinical reference in the available pilot.

## Limitations

1. Retrospective public-database design.
2. The RNA increment over clinical-only Cox was modest in TCGA nested cross-validation.
3. GEO platform heterogeneity limited absolute cutoff transferability and complicates calibration.
4. Cell-state analysis used curated signatures inferred from bulk RNA-seq rather than direct raw single-cell survival validation.
5. HPA evidence was qualitative and link based rather than a quantitative validation dataset.
6. CPTAC survival analysis was exploratory because only {evidence.cptac_deaths} deaths were available.
7. CPTAC quantified {evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} candidates rather than the complete candidate set.
8. The WSI pilot did not show stable improvement over clinical-only prediction.
9. No prospective clinical validation was performed.

## Conclusion

A frozen TCGA transcriptomic-clinical score reproduced risk ranking across four GEO cohorts and was associated with hypoxia, proliferative, EMT-like and stromal programs. HPA and exploratory CPTAC proteomics provided bounded orthogonal support, particularly for hypoxia and proliferation. These findings justify prospective, assay-harmonized validation, but not immediate clinical deployment or expansion into a larger multimodal predictor.

## Methods Skeleton

1. Study design and public data sources.
2. TCGA clinical and RNA-seq preprocessing.
3. Leakage-controlled nested cross-validation and frozen model fitting.
4. GEO expression harmonization and locked external evaluation.
5. Cell-state signature definitions and scoring.
6. HPA evidence collection.
7. PDC CPTAC-LUAD protein matrix processing and exploratory Cox analysis.
8. Exploratory WSI pilot.
9. Statistical analysis and multiplicity handling.
10. Reproducibility, code availability and data availability.

## Declarations To Complete

- Ethics statement for public de-identified data.
- Data availability with TCGA/GDC, GEO, HPA and PDC accession links.
- Code availability and software version statement.
- Funding, competing interests and author contributions.
- Reporting checklist selected for the target journal.

## Submission Boundary

This document is a manuscript skeleton, not a submission-ready final manuscript. Literature citations, author information, ethics wording, complete methods detail and journal-specific formatting remain to be completed.
"""


def write_manuscript_skeleton(
    root: str | Path = ".",
    *,
    create_docx: bool = True,
) -> dict[str, Path | bool]:
    paths = stage8_paths(root)
    paths["manuscript_dir"].mkdir(parents=True, exist_ok=True)
    evidence = load_stage8_evidence(root)
    markdown_path = paths["manuscript_dir"] / "SC_PROST_LUAD_manuscript_skeleton.md"
    markdown_path.write_text(manuscript_skeleton_markdown(evidence), encoding="utf-8")
    docx_path = paths["manuscript_dir"] / "SC_PROST_LUAD_manuscript_skeleton.docx"
    docx_created = False
    if create_docx:
        docx_created = markdown_to_docx(markdown_path, docx_path)
    return {
        "markdown": markdown_path,
        "docx": docx_path,
        "docx_created": docx_created,
    }


def markdown_to_docx(markdown_path: str | Path, output_path: str | Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return False
    document = Document()
    document.sections[0].top_margin = Inches(0.8)
    document.sections[0].bottom_margin = Inches(0.8)
    for raw_line in Path(markdown_path).read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            document.add_heading(_strip_markdown(heading.group(2)), level=min(len(heading.group(1)), 4))
        elif line.startswith("- "):
            document.add_paragraph(_strip_markdown(line[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(
                _strip_markdown(re.sub(r"^\d+\.\s+", "", line)),
                style="List Number",
            )
        else:
            document.add_paragraph(_strip_markdown(line))
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output.exists()


def _strip_markdown(text: str) -> str:
    return re.sub(r"[*`_]", "", text)


def build_figure_plan(root: str | Path = ".") -> pd.DataFrame:
    paths = stage8_paths(root)
    figures = paths["figures_dir"]
    rows = [
        {
            "figure_id": "Figure 1",
            "placement": "main",
            "title": "Study design and evidence workflow",
            "core_message": "Frozen TCGA model development was separated from GEO validation and from interpretation layers.",
            "panels": "A cohorts; B leakage-controlled model workflow; C locked GEO validation; D cell-state/HPA/CPTAC interpretation; E exploratory WSI boundary",
            "source_assets": "new schematic required",
            "source_tables": "stage2d_geo_dataset_summary.csv; stage5b_cptac_data_inventory.csv",
            "status": "design_required",
            "integrity_boundary": "Show WSI as exploratory and do not depict unperformed Stage 3 multimodal training.",
        },
        {
            "figure_id": "Figure 2",
            "placement": "main",
            "title": "TCGA model development and GEO external validation",
            "core_message": "RNA PCA_25 added a modest TCGA increment, while the frozen score retained discrimination across four GEO cohorts.",
            "panels": "A nested-CV C-index; B feature-space comparison; C external C-index; D external time-AUC",
            "source_assets": ";".join(
                str(figures / name)
                for name in (
                    "stage2c_nested_cv_cindex_boxplot.png",
                    "stage2c_feature_space_performance.png",
                    "stage2d_external_cindex_summary.png",
                    "stage2d_external_time_auc_summary.png",
                )
            ),
            "source_tables": "stage2c_feature_space_comparison.csv; stage2d_external_validation_performance.csv",
            "status": "source_panels_available",
            "integrity_boundary": "Do not claim large incremental value or neural-network superiority.",
        },
        {
            "figure_id": "Figure 3",
            "placement": "main",
            "title": "Risk stratification and survival associations across cohorts",
            "core_message": "Continuous risk scores were prognostic across cohorts, whereas the absolute TCGA cutoff was not portable.",
            "panels": "A TCGA OOF KM; B-D representative GEO KM; E external Cox forest; F cutoff transportability summary",
            "source_assets": ";".join(
                [str(figures / "stage2c_oof_km_best_model.png")]
                + [str(figures / f"stage2d_external_km_{cohort}.png") for cohort in ("GSE31210", "GSE50081", "GSE72094", "GSE68465")]
            ),
            "source_tables": "stage2c_oof_km_logrank.csv; stage2d_external_validation_cox.csv; stage2d_external_validation_performance.csv",
            "status": "forest_and_cutoff_panel_required",
            "integrity_boundary": "Label cohort-median KM as sensitivity analysis and state that the TCGA cutoff failed to transfer.",
        },
        {
            "figure_id": "Figure 4",
            "placement": "main",
            "title": "Cell-state interpretation of the frozen RNA risk score",
            "core_message": "High risk was associated with hypoxia, proliferation, EMT-like tumor and CAF programs, with external directional support.",
            "panels": "A TCGA correlations; B high-low score distributions; C adjusted Cox forest; D GEO consistency heatmap; E mechanism summary",
            "source_assets": ";".join(
                str(figures / name)
                for name in (
                    "stage4_tcga_risk_cell_state_correlation.png",
                    "stage4_tcga_cell_state_boxplots.png",
                    "stage4_tcga_cell_state_cox_forest.png",
                    "stage4_geo_cell_state_consistency_heatmap.png",
                    "stage4_mechanism_summary.png",
                )
            ),
            "source_tables": "stage4_tcga_cell_state_risk_correlation.csv; stage4_cell_state_external_consistency.csv",
            "status": "source_panels_available",
            "integrity_boundary": "Describe signature associations from bulk RNA, not direct single-cell outcome validation.",
        },
        {
            "figure_id": "Figure 5",
            "placement": "main",
            "title": "HPA and exploratory CPTAC protein-level evidence",
            "core_message": "Qualitative HPA evidence and CPTAC analyses prioritized hypoxia and proliferation without establishing causality.",
            "panels": "A HPA evidence overview; B CPTAC candidate abundance; C protein module scores; D exploratory Cox forest; E integrated evidence heatmap",
            "source_assets": ";".join(
                str(figures / name)
                for name in (
                    "stage5_hpa_evidence_summary.png",
                    "stage5b_cptac_candidate_boxplots.png",
                    "stage5b_cptac_module_score_heatmap.png",
                    "stage5b_cptac_protein_cox_forest.png",
                    "stage5b_integrated_hpa_cptac_heatmap.png",
                )
            ),
            "source_tables": "stage5_hpa_protein_evidence.csv; stage5b_integrated_hpa_cptac_evidence.csv; stage5b_cptac_protein_module_survival_cox.csv",
            "status": "source_panels_available",
            "integrity_boundary": "State 26 CPTAC deaths, 31/37 matching and exploratory protein support.",
        },
        {
            "figure_id": "Supplementary Figure S1",
            "placement": "supplementary",
            "title": "Exploratory WSI feasibility and neutral pilot findings",
            "core_message": "The pathology pipeline was technically feasible but did not improve on clinical-only prediction.",
            "panels": "A tissue mask; B patch grid; C attention heatmap; D model comparison; E overfitting diagnostics",
            "source_assets": ";".join(
                str(figures / name)
                for name in (
                    "stage6a_wsi_pilot_tissue_mask_example.png",
                    "stage6a_wsi_pilot_patch_grid_example.png",
                    "stage6a_wsi_pilot_attention_heatmap_example.png",
                    "stage6a_wsi_pilot_model_comparison.png",
                    "stage6a_mil_overfitting_heatmap.png",
                )
            ),
            "source_tables": "stage6a_wsi_pilot_model_performance.csv; stage6a_wsi_pilot_overfitting_diagnostics.csv",
            "status": "source_panels_available",
            "integrity_boundary": "Do not present WSI as a successful main prediction model.",
        },
    ]
    frame = pd.DataFrame(rows)
    frame["all_source_assets_exist"] = frame["source_assets"].map(_assets_exist)
    return frame


def _assets_exist(value: str) -> bool:
    if value == "new schematic required":
        return False
    return all(Path(item).exists() for item in value.split(";") if item)


def build_table_plan(root: str | Path = ".") -> pd.DataFrame:
    rows = [
        ("Table 1", "main", "TCGA and GEO cohort characteristics", "Cohort, platform, patients, deaths, censoring, clinical fields and gene coverage", "stage2_train_val_test_summary.csv; stage2d_geo_dataset_summary.csv; stage2d_external_validation_gene_missingness.csv"),
        ("Table 2", "main", "Model performance in TCGA and GEO cohorts", "Nested-CV C-index/AUC, external C-index/AUC and continuous-score Cox results", "stage2c_feature_space_comparison.csv; stage2d_external_validation_performance.csv; stage2d_external_validation_cox.csv"),
        ("Table 3", "main", "Cell-state associations with RNA risk and survival", "TCGA rho/q-value, high-low difference, adjusted Cox and external direction consistency", "stage4_tcga_cell_state_risk_correlation.csv; stage4_tcga_cell_state_group_comparison.csv; stage4_tcga_cell_state_multivariable_cox.csv; stage4_cell_state_external_consistency.csv"),
        ("Table 4", "main", "Integrated HPA and CPTAC evidence", "Candidate, mechanism layer, HPA status, CPTAC availability, stage association and exploratory survival statistics", "stage5b_integrated_hpa_cptac_evidence.csv; stage5b_cptac_protein_module_survival_cox.csv"),
        ("Table S1", "supplementary", "Full Stage 2D model performance", "All TCGA configurations, seeds and overfitting diagnostics", "stage2c_nested_cv_performance.csv; stage2c_overfitting_diagnostics.csv"),
        ("Table S2", "supplementary", "GEO preprocessing and missingness", "Platform, mapping, frozen-gene coverage and cohort-specific preprocessing", "stage2d_external_validation_gene_missingness.csv; stage2d_geo_dataset_summary.csv"),
        ("Table S3", "supplementary", "Cell-state signature definitions", "Signature names, categories, genes and expected directions", "stage4_cell_state_signature_definitions.csv"),
        ("Table S4", "supplementary", "Full cell-state association results", "All TCGA/GEO correlations, group comparisons and Cox analyses", "stage4_tcga_cell_state_cox.csv; stage4_geo_cell_state_risk_correlation.csv"),
        ("Table S5", "supplementary", "HPA candidate evidence and links", "All 37 candidate records, antibodies, reliability and pathology links", "stage5_hpa_protein_evidence.csv; stage5_hpa_ihc_links.csv"),
        ("Table S6", "supplementary", "CPTAC protein matching and exploratory survival", "31/37 matching, protein and module Cox results, and unavailable candidates", "stage5b_cptac_candidate_availability.csv; stage5b_cptac_protein_survival_cox.csv; stage5b_cptac_protein_module_survival_cox.csv"),
        ("Table S7", "supplementary", "Exploratory WSI pilot diagnostics", "Patch, feature, model and overfitting diagnostics", "stage6a_pathology_qc_summary.csv; stage6a_wsi_pilot_model_performance.csv; stage6a_wsi_pilot_overfitting_diagnostics.csv"),
    ]
    frame = pd.DataFrame(
        rows,
        columns=["table_id", "placement", "title", "content", "source_tables"],
    )
    frame["status"] = "assembly_required"
    frame["all_source_tables_exist"] = frame["source_tables"].map(
        lambda value: all(
            (stage8_paths(root)["tables_dir"] / item.strip()).exists()
            for item in value.split(";")
        )
    )
    return frame


def write_figure_table_outputs(root: str | Path = ".") -> dict[str, Path]:
    paths = stage8_paths(root)
    paths["tables_dir"].mkdir(parents=True, exist_ok=True)
    paths["reports_dir"].mkdir(parents=True, exist_ok=True)
    figure_plan = build_figure_plan(root)
    table_plan = build_table_plan(root)
    figure_path = paths["tables_dir"] / "stage8_figure_plan.csv"
    table_path = paths["tables_dir"] / "stage8_table_plan.csv"
    figure_plan.to_csv(figure_path, index=False)
    table_plan.to_csv(table_path, index=False)
    figure_legend_path = paths["reports_dir"] / "stage8_figure_legend_draft.md"
    table_legend_path = paths["reports_dir"] / "stage8_table_legend_draft.md"
    figure_legend_path.write_text(figure_legend_draft(), encoding="utf-8")
    table_legend_path.write_text(table_legend_draft(), encoding="utf-8")
    return {
        "figure_plan": figure_path,
        "table_plan": table_path,
        "figure_legends": figure_legend_path,
        "table_legends": table_legend_path,
    }


def results_narrative(evidence: Stage8Evidence) -> str:
    external_sentence = "; ".join(
        f"{row.cohort}, {row.c_index:.3f}"
        for row in evidence.external_performance.itertuples(index=False)
    )
    cox_sentence = "; ".join(
        f"{row.cohort}, HR {row.hazard_ratio_per_sd:.3f}, P={row.p_value:.3g}"
        for row in evidence.external_cox.itertuples(index=False)
    )
    core_states = evidence.cell_state_correlations.loc[
        evidence.cell_state_correlations["signature_name"].isin(
            [
                "hypoxia_tumor_cells",
                "proliferative_tumor_cells",
                "malignant_epithelial_cells",
                "emt_like_tumor_cells",
                "caf",
            ]
        )
    ]
    state_sentence = "; ".join(
        f"{row.signature_name}, rho={row.spearman_rho:.3f}, q={row.q_value_bh:.3g}"
        for row in core_states.itertuples(index=False)
    )
    return f"""# Stage 8 Results Narrative Draft

## Cohort construction and leakage-controlled model development

The TCGA-LUAD analysis included {evidence.tcga_patients} patients with usable overall survival and RNA-seq data, comprising {evidence.tcga_deaths} deaths and {evidence.tcga_censored} censored observations. All expression filtering, scaling, dimensionality reduction and model tuning were performed within training folds. Across repeated five-fold nested cross-validation, the RNA PCA_25 plus clinical Cox model achieved a mean C-index of {evidence.nested_combined_cindex:.3f} +/- {evidence.nested_combined_sd:.3f}, compared with {evidence.nested_clinical_cindex:.3f} +/- {evidence.nested_clinical_sd:.3f} for clinical Cox. Mean 1-, 3- and 5-year AUCs for the combined model were {evidence.nested_auc_1y:.3f}, {evidence.nested_auc_3y:.3f} and {evidence.nested_auc_5y:.3f}, respectively. Thus, the RNA representation provided a measurable but modest increment over the clinical reference. DeepSurv models showed substantial train-test gaps and did not support a neural-network performance advantage.

## Frozen external validation across four GEO cohorts

We next applied the frozen TCGA model to four independently processed GEO cohorts without refitting genes, scaling parameters, PCA loadings or Cox coefficients. The external cohorts comprised {evidence.external_total} patients in total. C-indices were {external_sentence}. The continuous risk score was associated with overall survival in all four cohorts ({cox_sentence}). These findings support reproducible risk ranking across microarray platforms. However, the absolute TCGA-derived median cutoff was not transportable: it failed to produce significant primary KM separation and assigned all samples to one group in two cohorts. Cohort-median KM analyses separated risk groups but were treated as sensitivity analyses rather than evidence for a transferable clinical threshold.

## Cell-state associations of the frozen risk score

To characterize the biological context of the risk score, we scored {evidence.signature_count} prespecified LUAD cell-state signatures in TCGA bulk RNA-seq. The principal high-risk associations were {state_sentence}. EMT-like, proliferative, hypoxic and CAF programs showed concordant directions across the four GEO cohorts. Dendritic-cell, B-cell and plasma-cell signatures were associated with the lower-risk direction. Several immune signatures, including M2 macrophage, Treg, exhausted CD8, cytotoxic CD8, M1 macrophage and NK programs, were mixed, contrary to expectation or externally unstable and were therefore excluded from the core mechanistic narrative. Cell-state scores with nominal age-, sex- and stage-adjusted Cox associations in TCGA included {", ".join(evidence.adjusted_cell_states)}. These results provide bulk RNA-based biological interpretation rather than direct single-cell prognostic validation.

## Orthogonal HPA and CPTAC protein evidence

We prespecified {evidence.hpa_candidate_count} marker genes spanning proliferation, hypoxia, EMT-like, CAF/matrix and dendritic/B/plasma contexts. All {evidence.hpa_supported_count} candidates had qualitative or linkable Human Protein Atlas evidence, although image-level lung-cancer staining was not available as a structured quantitative endpoint. We then analyzed the PDC CPTAC-LUAD discovery proteome, which included {evidence.cptac_samples} primary tumors and {evidence.cptac_proteins} quantified protein genes. Overall survival was usable for {evidence.cptac_usable_os} samples, with {evidence.cptac_deaths} deaths. Of the prespecified candidates, {evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} were quantified. LDHA, MKI67 and CDK1 showed nominally significant stage associations in the prespecified direction. No individual candidate protein remained significant after survival FDR correction. At the module level, the hypoxia score was associated with poorer survival (HR {evidence.cptac_hypoxia_hr:.3f} per SD, 95% CI {evidence.cptac_hypoxia_ci_low:.3f}-{evidence.cptac_hypoxia_ci_high:.3f}, P={evidence.cptac_hypoxia_p:.4f}, FDR={evidence.cptac_hypoxia_fdr:.4f}). The limited event count requires these protein findings to be interpreted as exploratory orthogonal support.

## Exploratory digital pathology feasibility

A separate {evidence.wsi_slides}-slide pilot established technical feasibility for OpenSlide processing, tissue patch extraction, CUDA ResNet50 feature extraction and MIL training. The clinical-only test C-index was {evidence.wsi_clinical_cindex:.3f}, whereas the best pathology-only and fusion test C-indices were {evidence.wsi_pathology_cindex:.3f} and {evidence.wsi_fusion_cindex:.3f}. Pathology and fusion models therefore did not show stable improvement over the clinical reference. This analysis should remain supplementary as a neutral feasibility and limitation result.
"""


def discussion_points(evidence: Stage8Evidence) -> str:
    return f"""# Stage 8 Discussion Points

## Central Advance

- The strongest contribution is external reproducibility of a frozen transcriptomic-clinical risk ranking across four independent GEO cohorts.
- The TCGA RNA increment was modest: C-index {evidence.nested_combined_cindex:.3f} versus {evidence.nested_clinical_cindex:.3f} for clinical Cox.
- The biological contribution is a convergent association with hypoxia, proliferation, EMT-like tumor and CAF programs.

## Interpretation

- Emphasize continuous risk ranking rather than a transferable absolute threshold.
- Discuss the failure of the TCGA median cutoff as evidence of platform and calibration shift.
- Treat cell-state scores as bulk RNA-derived associations informed by curated single-cell literature.
- Position LDHA, MKI67 and CDK1 as prioritized exploratory protein candidates.
- Highlight the hypoxia module survival result while stating that CPTAC had only {evidence.cptac_deaths} deaths.

## Negative And Neutral Results

- DeepSurv overfitting argues for the simpler Cox model.
- RNA added only a small nested-CV increment over clinical predictors.
- WSI and clinical-pathology fusion did not improve on clinical-only prediction in the pilot.
- These neutral findings improve credibility and define where additional complexity was not justified.

## Translational Implication

- The present score is suitable for research stratification and hypothesis generation.
- Future work should prioritize prospective sampling, assay harmonization, prespecified thresholds, calibration and decision-utility analysis.
- Protein experiments should focus on the hypoxia/proliferation axis rather than claiming broad confirmation of every cell-state layer.

## Reviewer-Facing Risks

- Incremental value over clinical Cox may be considered small.
- Absolute cutoff transport failed across GEO platforms.
- External cohorts use heterogeneous microarray platforms and clinical annotations.
- The cell-state layer does not use raw single-cell outcome data.
- CPTAC event counts limit survival inference.
- WSI results are neutral and should not be used to advertise a multimodal predictor.
"""


def limitations_text(evidence: Stage8Evidence) -> str:
    return f"""# Stage 8 Limitations

1. This study used a retrospective public-database design, with cohort-specific differences in treatment, follow-up and annotation.
2. The RNA PCA_25 representation produced only a modest increment over clinical-only Cox in TCGA nested cross-validation.
3. GEO platform heterogeneity affected frozen-gene coverage, calibration and absolute cutoff transferability.
4. Cell-state analyses used curated signatures inferred from bulk RNA-seq rather than raw single-cell data with linked survival outcomes.
5. Human Protein Atlas evidence was qualitative and link based; structured quantitative lung-cancer IHC measurements were not used.
6. CPTAC analyses were exploratory because only {evidence.cptac_deaths} deaths were available among {evidence.cptac_usable_os} samples with usable overall survival.
7. CPTAC quantified {evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} prespecified candidates rather than the complete candidate set.
8. The WSI pilot did not show stable improvement over clinical-only prediction and was not expanded into a full pathology modeling study.
9. No prospective clinical validation, prespecified clinical threshold evaluation or deployment study was performed.
"""


def write_results_outputs(root: str | Path = ".") -> dict[str, Path]:
    paths = stage8_paths(root)
    paths["reports_dir"].mkdir(parents=True, exist_ok=True)
    evidence = load_stage8_evidence(root)
    outputs = {
        "results": paths["reports_dir"] / "stage8_results_narrative_draft.md",
        "discussion": paths["reports_dir"] / "stage8_discussion_points.md",
        "limitations": paths["reports_dir"] / "stage8_limitations.md",
    }
    outputs["results"].write_text(results_narrative(evidence), encoding="utf-8")
    outputs["discussion"].write_text(discussion_points(evidence), encoding="utf-8")
    outputs["limitations"].write_text(limitations_text(evidence), encoding="utf-8")
    return outputs


def figure_legend_draft() -> str:
    return """# Stage 8 Figure Legend Draft

## Figure 1. Study design and evidence workflow

Overview of the SC-PROST-LUAD study. TCGA-LUAD clinical and RNA-seq data were used for leakage-controlled model development. The frozen PCA_25 plus clinical Cox model was applied without refitting to four GEO cohorts. Curated cell-state signatures, HPA records and PDC CPTAC-LUAD proteomics were used as interpretation layers. WSI analysis was restricted to an exploratory feasibility pilot.

## Figure 2. TCGA model development and GEO external validation

Nested cross-validation comparison of clinical and RNA feature spaces, followed by locked external evaluation in GSE31210, GSE50081, GSE72094 and GSE68465. Error bars show between-fold or cohort variability as specified in each panel. External values were generated without refitting the frozen TCGA model.

## Figure 3. Risk stratification and survival associations across cohorts

Out-of-fold TCGA survival stratification and GEO survival analyses. Continuous-score Cox results are primary external association results. Cohort-median KM curves are sensitivity analyses because the absolute TCGA median cutoff did not transfer reliably across platforms.

## Figure 4. Cell-state interpretation of the frozen RNA risk score

Associations between the frozen risk score and 16 curated LUAD cell-state signatures in TCGA, high-versus-low risk comparisons, adjusted Cox analyses and directional consistency across four GEO cohorts. Scores were inferred from bulk expression and represent biological associations rather than direct single-cell outcome validation.

## Figure 5. HPA and exploratory CPTAC protein-level evidence

Qualitative HPA evidence and quantitative PDC CPTAC-LUAD analyses for prespecified candidate proteins. CPTAC analyses included 110 primary tumors and 26 deaths. Directional stage associations and exploratory survival results are shown; the evidence is orthogonal and does not establish causality.

## Supplementary Figure S1. Exploratory WSI feasibility and neutral pilot findings

Representative tissue segmentation, patch extraction, feature visualization, model comparison and overfitting diagnostics from the 100-slide WSI pilot. Pathology and fusion models did not show stable improvement over the clinical-only reference.
"""


def table_legend_draft() -> str:
    return """# Stage 8 Table Legend Draft

## Table 1. TCGA and external GEO cohort characteristics

Clinical, survival, platform and expression-coverage characteristics for TCGA-LUAD and the four GEO validation cohorts.

## Table 2. Model performance in TCGA and external GEO cohorts

Leakage-controlled nested-CV performance in TCGA and locked external discrimination, time-AUC and continuous-score Cox associations in GEO.

## Table 3. Cell-state associations with RNA risk score and survival

TCGA risk-score correlations, risk-group comparisons, adjusted Cox results and directional consistency across GEO cohorts for 16 curated signatures.

## Table 4. Integrated HPA and CPTAC protein evidence

Prespecified candidate proteins, qualitative HPA evidence, CPTAC availability, direction-consistent stage associations and exploratory survival statistics.

## Supplementary Tables S1-S7

Full model results, GEO preprocessing, signature definitions, cell-state analyses, HPA records, CPTAC matching and survival analyses, and exploratory WSI diagnostics.
"""


def journal_targeting_report() -> str:
    today = date.today().isoformat()
    return f"""# Stage 8 Journal Targeting Report

Retrieved from official journal pages on {today}. Journal fit should be rechecked immediately before submission.

## Recommended Order

### 1. BMC Medical Genomics - Best Current Fit

- Official scope: genomic research in human health and disease, including proteomics, cancer genomics, precision medicine and bioinformatics methods.
- Fit: strong for a scientifically valid public-data genomics study with external validation and bounded protein interpretation.
- Positioning: emphasize frozen external validation, reproducibility and integrated genomic/proteomic interpretation.
- Main risk: reviewers may request more formal clinical utility analysis and clearer comparison with published LUAD signatures.
- Official scope: https://link.springer.com/journal/12920/aims-and-scope

### 2. Journal of Translational Medicine - Ambitious Translational Option

- Official scope includes disease biomarkers, medical bioinformatics, computational modelling, data-driven clinical decision processes, molecular pathology and translational cancer biology.
- Fit: reasonable if the manuscript foregrounds external validation and biological translation rather than algorithmic novelty.
- Main risk: modest RNA increment, cutoff non-transferability and absence of prospective validation weaken the clinical-translation case.
- Official scope: https://link.springer.com/journal/12967/aims-and-scope

### 3. npj Precision Oncology - Stretch Target

- Official scope explicitly includes biomarker studies with biological and independent external validation, tumor microenvironment research and externally validated computational oncology with clear clinical potential.
- Fit: external validation and multi-layer interpretation are aligned.
- Main risk: the current evidence does not yet establish strong clinical potential, calibrated decision utility or prospective validation. A presubmission enquiry is advisable.
- Official scope: https://www.nature.com/npjprecisiononcology/aims

### 4. Scientific Reports - Broad-Scope Fallback

- Official scope covers original research across medicine, biomedical sciences and engineering, with emphasis on scientific robustness.
- Fit: appropriate for a reproducible computational oncology study when claims remain bounded.
- Main risk: the manuscript still needs a clear advance beyond another public-database prognostic signature study.
- Official information: https://www.nature.com/srep/about

## Submission Strategy

1. Complete the manuscript for BMC Medical Genomics formatting first.
2. Before submission, add a structured comparison with representative externally validated LUAD prognostic models.
3. Keep WSI in the supplement.
4. Lead with locked external validation and cutoff-transfer limitations, not with multimodal complexity.
5. Present HPA and CPTAC as orthogonal interpretation, with CPTAC explicitly labeled exploratory.

## Not Recommended Without New Evidence

- High-impact clinical oncology journals that require prospective utility or treatment-decision relevance.
- Digital pathology journals if the central claim depends on pathology improving prediction.
- Single-cell journals because this study does not analyze a raw single-cell cohort with linked survival.
"""


def write_journal_targeting_report(root: str | Path = ".") -> Path:
    paths = stage8_paths(root)
    paths["reports_dir"].mkdir(parents=True, exist_ok=True)
    output = paths["reports_dir"] / "stage8_journal_targeting_report.md"
    output.write_text(journal_targeting_report(), encoding="utf-8")
    return output


def evidence_audit_markdown(root: str | Path = ".") -> str:
    evidence = load_stage8_evidence(root)
    required_files = [
        "outputs/manuscript/SC_PROST_LUAD_manuscript_skeleton.md",
        "outputs/manuscript/SC_PROST_LUAD_manuscript_skeleton.docx",
        "outputs/tables/stage8_figure_plan.csv",
        "outputs/tables/stage8_table_plan.csv",
        "outputs/reports/stage8_results_narrative_draft.md",
        "outputs/reports/stage8_discussion_points.md",
        "outputs/reports/stage8_limitations.md",
        "outputs/reports/stage8_figure_legend_draft.md",
        "outputs/reports/stage8_table_legend_draft.md",
        "outputs/reports/stage8_journal_targeting_report.md",
    ]
    root_path = Path(root).resolve()
    missing = [item for item in required_files if not (root_path / item).exists()]
    claim_rows = [
        ("Frozen score externally discriminates risk", "Four GEO C-indices 0.624-0.700 and significant continuous-score Cox results", "supported"),
        ("RNA materially improves clinical prediction", f"Nested C-index {evidence.nested_combined_cindex:.3f} versus {evidence.nested_clinical_cindex:.3f}", "must be qualified as modest"),
        ("A universal high-risk cutoff transfers", "TCGA cutoff failed to separate cohorts and collapsed two cohorts into one group", "not supported"),
        ("Hypoxia/proliferation are associated with higher risk", "TCGA/GEO signature consistency plus LDHA/MKI67/CDK1 and hypoxia module evidence", "supported as association"),
        ("All protein candidates have quantitative support", f"CPTAC matched {evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total}; only {len(evidence.cptac_supported_genes)} met the stage-direction support rule", "not supported"),
        ("WSI improves prediction", f"Clinical {evidence.wsi_clinical_cindex:.3f}; pathology {evidence.wsi_pathology_cindex:.3f}; fusion {evidence.wsi_fusion_cindex:.3f}", "not supported"),
        ("Model is ready for clinical deployment", "Retrospective cohorts, cutoff shift and no prospective validation", "not supported"),
    ]
    claim_table = "\n".join(
        f"| {claim} | {support} | {status} |"
        for claim, support, status in claim_rows
    )
    return f"""# Stage 8 Final Evidence Audit

## Scope

This audit checks claim-evidence alignment for the Stage 8 manuscript skeleton. No model was retrained and no Stage 3 or Stage 6B/6C analysis was initiated.

## Output Completeness

- Required outputs missing: {", ".join(missing) if missing else "none"}.
- DOCX dependency available: manuscript DOCX generated if present in the list above.

## Claim-Evidence Map

| Proposed claim | Direct evidence | Audit status |
| --- | --- | --- |
{claim_table}

## Evidence Hierarchy

1. Primary predictive evidence: leakage-controlled TCGA nested CV and locked GEO external validation.
2. Primary biological interpretation: prespecified bulk RNA cell-state associations with external directional checks.
3. Orthogonal support: qualitative HPA links and exploratory CPTAC protein analyses.
4. Supplementary neutral evidence: technically feasible WSI pilot without stable predictive gain.

## Required Boundaries

- Use continuous external risk ranking as the central validation claim.
- State that the RNA increment over clinical Cox was modest.
- State that the TCGA cutoff did not transfer reliably.
- Describe cell states as signature-based associations inferred from bulk RNA.
- Describe HPA as qualitative/link evidence.
- Describe CPTAC as exploratory because only {evidence.cptac_deaths} deaths were available and {evidence.cptac_candidates_matched}/{evidence.cptac_candidate_total} candidates were matched.
- Keep WSI outside the main predictive claim.
- State that prospective clinical validation remains absent.

## Reviewer Risk Assessment

- Contribution risk: the model itself is simple and its incremental value is small; novelty rests on locked external validation and a carefully bounded interpretation chain.
- Validation risk: external ranking is reproducible, but calibration and threshold portability are incomplete.
- Mechanism risk: signatures and proteins support association, not causation.
- Clinical risk: no prospective utility or decision-impact study is available.

## Final Recommendation

Proceed to manuscript drafting with BMC Medical Genomics as the most realistic initial target. Do not describe the current artifact as submission-ready until literature citations, comparison with published LUAD signatures, complete methods, ethics wording, authorship and journal formatting are added.
"""


def write_final_evidence_audit(root: str | Path = ".") -> dict[str, Path | int]:
    paths = stage8_paths(root)
    paths["reports_dir"].mkdir(parents=True, exist_ok=True)
    output = paths["reports_dir"] / "stage8_final_evidence_audit.md"
    output.write_text(evidence_audit_markdown(root), encoding="utf-8")
    scan_paths = [
        paths["manuscript_dir"] / "SC_PROST_LUAD_manuscript_skeleton.md",
        paths["reports_dir"] / "stage8_results_narrative_draft.md",
        paths["reports_dir"] / "stage8_discussion_points.md",
        paths["reports_dir"] / "stage8_limitations.md",
    ]
    forbidden = scan_forbidden_overclaims(scan_paths)
    return {"report": output, "forbidden_phrase_hits": len(forbidden)}


def scan_forbidden_overclaims(paths: list[Path]) -> list[dict[str, str]]:
    patterns = [
        "causal confirmation",
        "clinically ready tool",
        "definitive proteomic validation",
        "single-cell validated survival model",
        "digital pathology improved prediction",
        "multimodal deep learning superiority",
        "was prospectively validated",
        "prospectively validated model",
        "all candidates validated by cptac",
    ]
    hits: list[dict[str, str]] = []
    for path in paths:
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8").lower()
        for pattern in patterns:
            if pattern in text:
                hits.append({"file": str(path), "phrase": pattern})
    return hits


def append_stage8_audit(root: str | Path, message: str) -> Path:
    path = stage8_paths(root)["audit"]
    with path.open("a", encoding="utf-8") as handle:
        handle.write(message)
    return path
